"""Generador de los 5 documentos .docx para entregar en Azure.

Genera:
  1_SRS.docx                       — Especificación de Requisitos del Software
  2_Diagrama_Casos_de_Uso.docx     — Diagrama UML de casos de uso
  3_Especificacion_Casos_de_Uso.docx — Detalle de >= 5 casos de uso
  4_Diagramas_Comportamiento.docx  — Secuencia, actividad, estado
  5_Diagramas_Estructurales.docx   — Clases, componentes, despliegue

Las imágenes se generan con generar_diagramas.py (matplotlib).
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = Path(__file__).parent
IMG = BASE / "imagenes"

# Paleta corporativa Distrito Urbano
BRAND = RGBColor(0x1F, 0x7A, 0x5C)
BRAND_DARK = RGBColor(0x13, 0x4E, 0x3A)
ACCENT = RGBColor(0xF5, 0x9E, 0x0B)
GRAY = RGBColor(0x4C, 0x59, 0x60)
TEXT = RGBColor(0x17, 0x20, 0x26)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


# =========================================================================
# Helpers comunes (estilo Word)
# =========================================================================

def set_doc_lang(doc, lang="es-CO"):
    """Forzar el idioma del documento a español (Colombia) para que Word no
    subraye en rojo palabras correctas."""
    styles_element = doc.styles.element
    for rpr_default in styles_element.iter(qn("w:rPrDefault")):
        rpr = rpr_default.find(qn("w:rPr"))
        if rpr is None:
            rpr = OxmlElement("w:rPr")
            rpr_default.insert(0, rpr)
        # Quitar cualquier w:lang previo dentro de rPr
        for lang_el in list(rpr.findall(qn("w:lang"))):
            rpr.remove(lang_el)
        lang_el = OxmlElement("w:lang")
        lang_el.set(qn("w:val"), lang)
        lang_el.set(qn("w:eastAsia"), lang)
        lang_el.set(qn("w:bidi"), "ar-SA")
        rpr.append(lang_el)


def set_margins(doc, top=2, bottom=2, left=2.2, right=2.2):
    for section in doc.sections:
        section.top_margin = Cm(top)
        section.bottom_margin = Cm(bottom)
        section.left_margin = Cm(left)
        section.right_margin = Cm(right)


def set_run(run, *, bold=False, italic=False, size=11, color=TEXT, font="Calibri"):
    run.font.name = font
    r = run._element
    r.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def add_para(doc, text, *, bold=False, italic=False, size=11, color=TEXT,
             align=WD_ALIGN_PARAGRAPH.JUSTIFY, before=4, after=4):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    run = p.add_run(text)
    set_run(run, bold=bold, italic=italic, size=size, color=color)
    return p


def add_title(doc, text, color=BRAND_DARK, size=22):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run(run, bold=True, size=size, color=color)
    return p


def add_h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run(run, bold=True, size=16, color=BRAND)
    # Línea inferior
    p_pr = p._p.get_or_add_pPr()
    border = OxmlElement("w:pBdr")
    b = OxmlElement("w:bottom")
    b.set(qn("w:val"), "single"); b.set(qn("w:sz"), "8")
    b.set(qn("w:color"), "1F7A5C")
    border.append(b)
    p_pr.append(border)
    return p


def add_h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run(run, bold=True, size=13, color=BRAND_DARK)
    return p


def add_h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    set_run(run, bold=True, size=11.5, color=GRAY)
    return p


def add_bullet(doc, text, level=0, bold=False):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.6 + level * 0.6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    set_run(run, size=10.5, bold=bold)
    return p


def add_image(doc, path, width_cm=15.5, caption=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Cm(width_cm))
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(f"Figura. {caption}")
        set_run(run, italic=True, size=9.5, color=GRAY)


def add_table_header(doc, headers, widths_cm=None):
    """Crea una tabla con encabezado coloreado en BRAND y devuelve la tabla."""
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        if widths_cm:
            cell.width = Cm(widths_cm[i])
        # Color de fondo
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "1F7A5C")
        tc_pr.append(shd)
        # Texto en blanco
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_run(run, bold=True, size=10, color=WHITE)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    return t


def add_table_row(table, cells, sizes=10):
    row = table.add_row()
    if isinstance(sizes, int):
        sizes = [sizes] * len(cells)
    for i, txt in enumerate(cells):
        c = row.cells[i]
        c.text = ""
        p = c.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(str(txt))
        set_run(run, size=sizes[i])
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    return row


def add_cover(doc, *, entregable, materia, profesor, institucion):
    """Portada estándar para todos los entregables."""
    # Espacio superior
    for _ in range(3):
        doc.add_paragraph()
    # Marca DU
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("DISTRITO URBANO")
    set_run(run, bold=True, size=28, color=BRAND)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Tienda Digital")
    set_run(run, italic=True, size=14, color=GRAY)

    doc.add_paragraph()
    doc.add_paragraph()

    # Título del entregable
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(entregable)
    set_run(run, bold=True, size=22, color=BRAND_DARK)

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    # Bloque institucional
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(institucion); set_run(run, bold=True, size=12, color=TEXT)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(materia); set_run(run, size=11.5, color=TEXT)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Docente: {profesor}"); set_run(run, italic=True, size=10.5, color=GRAY)

    doc.add_paragraph()
    doc.add_paragraph()

    # Integrantes
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Integrantes")
    set_run(run, bold=True, size=11.5, color=BRAND)

    integrantes = [
        "Daniel Villamizar — Frontend",
        "Santiago Pérez — Backend",
        "Tomás Urieles — Scrum Master / Infraestructura",
    ]
    for i in integrantes:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(i)
        set_run(run, size=10.5, color=TEXT)

    # Fecha
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Mayo de 2026")
    set_run(run, italic=True, size=10.5, color=GRAY)

    doc.add_page_break()


COVER_KW = dict(
    materia="Ingeniería de Software",
    profesor="(Docente del curso)",
    institucion="Universidad Autónoma de Bucaramanga (UNAB)",
)


# =========================================================================
# DATA — Requisitos funcionales y no funcionales
# =========================================================================

RF = [
    ("RF-01", "Gestión de cuentas, autenticación y control de acceso",
     "El sistema debe permitir la gestión de cuentas, autenticación y control de acceso de clientes y administradores.",
     "Yo como Product Owner, para controlar el acceso seguro de clientes y administradores a la tienda digital, "
     "necesito que el sistema permita crear cuentas, iniciar sesión, cerrar sesión, actualizar perfiles administrativos "
     "y aplicar control de acceso basado en roles.",
     [
         ("Crear formulario de registro de cliente",
          "Diseñar e implementar el formulario para registrar nombre, correo electrónico, teléfono y contraseña."),
         ("Implementar inicio y cierre de sesión",
          "Crear la lógica de autenticación para clientes y administradores, incluyendo validación de credenciales y cierre seguro de sesión."),
         ("Implementar control de acceso por rol",
          "Restringir vistas y operaciones según el rol del usuario, diferenciando cliente y administrador."),
         ("Crear módulo de perfil administrativo",
          "Permitir que el administrador actualice su perfil, credenciales y datos de contacto."),
         ("Validar seguridad básica de credenciales",
          "Verificar datos obligatorios, correo único, formato de contraseña y acceso restringido a módulos protegidos."),
     ]),
    ("RF-02", "Vitrina pública con identidad visual",
     "El sistema debe mostrar la tienda digital, categorías y productos activos con identidad visual propia de la empresa.",
     "Yo como Product Owner, para ofrecer una experiencia visual clara y coherente con la marca de la empresa, "
     "necesito que el sistema muestre la tienda digital con logo, colores, mensajes informativos, categorías y productos activos.",
     [
         ("Diseñar interfaz principal de la tienda",
          "Crear la pantalla principal de la tienda digital con estructura visual, navegación y espacios para identidad corporativa."),
         ("Mostrar identidad visual de la empresa",
          "Integrar logo, colores corporativos, banner principal y mensajes informativos visibles para los clientes."),
         ("Listar categorías activas",
          "Consultar y mostrar las categorías disponibles dentro del catálogo."),
         ("Listar productos activos",
          "Mostrar productos con nombre, imágenes, precio y disponibilidad."),
         ("Aplicar estado de publicación",
          "Garantizar que solo se visualicen en la tienda los productos y categorías activas."),
     ]),
    ("RF-03", "Búsqueda, filtros y ficha de producto",
     "El sistema debe permitir buscar, filtrar y consultar el detalle de productos con sus variantes.",
     "Yo como Product Owner, para facilitar que el cliente encuentre y revise productos antes de comprar, "
     "necesito que el sistema permita buscar, filtrar y consultar la ficha completa de productos con descripción, "
     "galería, variantes, precio, stock visible y valoración promedio.",
     [
         ("Implementar buscador de productos",
          "Permitir búsqueda de productos por texto ingresado por el cliente."),
         ("Implementar filtros de catálogo",
          "Permitir filtrar productos por categoría, rango de precio, color, talla, tamaño o disponibilidad."),
         ("Crear ficha de producto",
          "Mostrar descripción, galería de imágenes, precio, stock visible y valoración promedio."),
         ("Mostrar variantes del producto",
          "Presentar opciones de color, talla, tamaño u otros atributos configurables."),
         ("Validar selección de variante",
          "Exigir que el cliente seleccione una variante específica antes de agregar el producto al carrito."),
     ]),
    ("RF-04", "Carrito de compras y captura de datos de checkout",
     "El sistema debe permitir gestionar el carrito de compras, validar stock y capturar datos para finalizar el pedido.",
     "Yo como Product Owner, para permitir que el cliente prepare correctamente su compra antes del pago, "
     "necesito que el sistema permita agregar productos al carrito, modificar cantidades, eliminar ítems, "
     "validar stock y capturar datos de entrega, facturación y contacto.",
     [
         ("Implementar carrito de compras",
          "Permitir agregar productos seleccionados al carrito con su respectiva variante."),
         ("Modificar cantidades del carrito",
          "Permitir que el cliente aumente o disminuya cantidades antes de confirmar la compra."),
         ("Eliminar productos del carrito",
          "Permitir retirar ítems del carrito y actualizar el resumen de compra."),
         ("Validar disponibilidad de stock",
          "Verificar que la variante seleccionada tenga stock suficiente antes de iniciar el checkout."),
         ("Capturar datos de checkout",
          "Registrar información de entrega, facturación y contacto necesaria para finalizar el pedido."),
     ]),
    ("RF-05", "Cálculo, pasarela y creación del pedido",
     "El sistema debe calcular valores de compra, integrar pasarela de pago y crear pedidos con estado actualizado.",
     "Yo como Product Owner, para procesar correctamente la compra del cliente, necesito que el sistema calcule "
     "subtotal, costos adicionales, descuentos, total a pagar, integre una pasarela de pago externa y cree el pedido "
     "con un identificador único y estado según el resultado del pago.",
     [
         ("Calcular valores de compra",
          "Implementar cálculo de subtotal, costos adicionales, descuentos aplicables y total a pagar."),
         ("Mostrar resumen de pago",
          "Presentar al cliente el detalle económico de la compra antes de pagar."),
         ("Integrar pasarela de pago externa",
          "Conectar el sistema con una pasarela que permita autorizar, rechazar o dejar pendiente la transacción."),
         ("Crear pedido con identificador único",
          "Registrar el pedido generado con un código único y los datos de compra asociados."),
         ("Validar respuesta de la pasarela",
          "Impedir la confirmación del pedido si el pago no cuenta con respuesta válida."),
     ]),
    ("RF-06", "Consulta, actualización y notificación de pedidos",
     "El sistema debe permitir consultar, actualizar y notificar el estado de los pedidos.",
     "Yo como Product Owner, para permitir el seguimiento completo de las compras realizadas, necesito que el sistema "
     "permita al cliente consultar su historial y detalle de pedidos, manejar estados del pedido, permitir actualizaciones "
     "operativas por parte del administrador y notificar cambios relevantes.",
     [
         ("Crear historial de pedidos del cliente",
          "Permitir que el cliente consulte los pedidos realizados desde su cuenta."),
         ("Mostrar detalle del pedido",
          "Presentar productos, cantidades, valores, datos de entrega y estado actual del pedido."),
         ("Implementar estados del pedido",
          "Manejar estados como pagado, en preparación, enviado, entregado y cancelado."),
         ("Permitir actualización administrativa del pedido",
          "Habilitar al administrador para actualizar el estado operativo conforme avance el proceso logístico."),
         ("Notificar cambios relevantes",
          "Informar al cliente sobre cambios importantes del pedido y resultado del pago."),
     ]),
    ("RF-07", "Gestión de catálogo, variantes e inventario",
     "El sistema debe permitir la gestión administrativa de categorías, productos, variantes e inventario.",
     "Yo como Product Owner, para administrar correctamente el catálogo y el inventario de la tienda, "
     "necesito que el sistema permita crear, consultar, editar, archivar y reactivar categorías y productos, "
     "definir variantes, asignar SKU, costo, precio y stock, registrar movimientos de inventario y generar alertas de stock mínimo.",
     [
         ("Implementar gestión de categorías",
          "Permitir crear, consultar, editar, archivar y reactivar categorías de producto."),
         ("Implementar gestión de productos",
          "Permitir crear, consultar, editar, archivar y reactivar productos."),
         ("Asociar información comercial al producto",
          "Permitir cargar imágenes, descripciones extensas, precios y estado de publicación."),
         ("Gestionar variantes de producto",
          "Permitir definir variantes por color, talla, tamaño u otros atributos, asignando SKU, costo, precio y stock independiente."),
         ("Registrar movimientos y alertas de inventario",
          "Registrar entradas, salidas, ajustes y reservas con motivo y usuario responsable, generando alertas cuando se alcance el umbral mínimo de stock."),
     ]),
    ("RF-08", "Información administrativa, financiera y reportes",
     "El sistema debe permitir gestionar información administrativa, financiera y reportes del negocio.",
     "Yo como Product Owner, para controlar la operación administrativa y financiera del negocio, "
     "necesito que el sistema permita registrar empleados, costos operativos, gastos generales, calcular indicadores financieros, "
     "mostrar dashboards y exportar reportes administrativos y financieros.",
     [
         ("Registrar empleados",
          "Permitir registrar empleados con datos básicos, cargo, salario y estado laboral."),
         ("Registrar costos y gastos",
          "Permitir registrar costos operativos, gastos generales y observaciones asociadas al negocio."),
         ("Calcular indicadores financieros",
          "Calcular ventas brutas, COGS, costos operativos, nómina y utilidad neta por período."),
         ("Crear dashboards administrativos",
          "Mostrar ventas por período, productos más vendidos, rotación de inventario y resumen financiero."),
         ("Exportar reportes",
          "Permitir exportar reportes administrativos y financieros en formatos comunes como PDF o CSV."),
     ]),
    ("RF-09", "Configuración de la tienda, mensajes y reseñas",
     "El sistema debe permitir la configuración general de la tienda, mensajes informativos, parámetros globales y reseñas de productos.",
     "Yo como Product Owner, para administrar la configuración visible y operativa de la tienda digital, "
     "necesito que el sistema permita configurar datos generales, administrar mensajes informativos, definir parámetros globales "
     "del negocio y permitir reseñas únicamente cuando exista una compra completada y entregada.",
     [
         ("Configurar datos generales de la tienda",
          "Permitir configurar nombre comercial, logo, colores, banner principal y datos de contacto."),
         ("Administrar mensajes informativos",
          "Permitir crear, editar y publicar mensajes visibles para clientes, como horarios, políticas y avisos temporales."),
         ("Definir parámetros globales",
          "Permitir configurar moneda, umbral de stock y estados operativos habilitados."),
         ("Implementar reseñas y valoraciones",
          "Permitir que el cliente valore y reseñe productos comprados."),
         ("Validar compra entregada para reseñas",
          "Restringir las reseñas únicamente a productos asociados a una compra completada y entregada."),
     ]),
]


RNF = [
    ("RNF-01", "Rendimiento, disponibilidad y concurrencia",
     "El sistema debe garantizar rendimiento, disponibilidad y capacidad concurrente bajo carga normal.",
     "Yo como Product Owner, para asegurar que la tienda digital funcione de forma estable y rápida durante su operación normal, "
     "necesito que el sistema responda en tiempos adecuados, soporte usuarios concurrentes y mantenga una disponibilidad mensual aceptable.",
     [
         ("Validar tiempo de respuesta del catálogo y panel",
          "Verificar que el 95% de las operaciones de consulta de catálogo y panel respondan en 3 segundos o menos bajo carga normal."),
         ("Realizar pruebas de concurrencia de clientes",
          "Comprobar que la solución soporte al menos 300 clientes concurrentes."),
         ("Realizar pruebas de concurrencia administrativa",
          "Comprobar que la solución soporte al menos 20 sesiones administrativas concurrentes."),
         ("Implementar monitoreo de disponibilidad",
          "Medir que la plataforma mantenga una disponibilidad mensual de 99,5% o superior, excluyendo mantenimientos programados."),
         ("Optimizar consultas y recursos críticos",
          "Mejorar consultas, endpoints y carga de información para cumplir los tiempos definidos."),
     ]),
    ("RNF-02", "Seguridad, privacidad e integridad",
     "El sistema debe garantizar seguridad, privacidad e integridad de la información y las transacciones.",
     "Yo como Product Owner, para proteger los datos personales, credenciales, operaciones administrativas y transacciones de compra, "
     "necesito que el sistema aplique comunicación segura, almacenamiento protegido de contraseñas, control de acceso, "
     "mínimo privilegio e integridad transaccional.",
     [
         ("Implementar comunicación HTTPS/TLS",
          "Garantizar que toda comunicación entre frontend, backend y servicios externos se realice mediante HTTPS/TLS."),
         ("Proteger contraseñas y sesiones",
          "Almacenar contraseñas mediante hash robusto y configurar expiración de sesiones por inactividad."),
         ("Restringir funciones administrativas",
          "Asegurar que las funciones administrativas solo estén disponibles para usuarios autenticados con rol administrador."),
         ("Aplicar principio de mínimo privilegio",
          "Controlar el acceso a datos personales y transaccionales según permisos definidos."),
         ("Garantizar integridad transaccional",
          "Evitar doble descuento de inventario y evitar marcar pedidos como pagados sin confirmación válida de la pasarela."),
     ]),
    ("RNF-03", "Usabilidad, mantenibilidad, respaldo y trazabilidad",
     "El sistema debe garantizar usabilidad, compatibilidad, mantenibilidad, respaldo, trazabilidad y observabilidad.",
     "Yo como Product Owner, para asegurar que la solución sea usable, mantenible, recuperable y fácil de diagnosticar, "
     "necesito que el sistema sea compatible con navegadores modernos, tenga interfaz accesible, arquitectura modular, "
     "respaldo de información, trazabilidad de operaciones críticas y registros suficientes para diagnosticar fallos.",
     [
         ("Validar usabilidad y accesibilidad",
          "Verificar que un usuario nuevo pueda completar su primera compra en 5 minutos o menos y que la interfaz mantenga contraste, "
          "navegación consistente y ampliación de texto hasta 200%."),
         ("Validar compatibilidad multidispositivo",
          "Comprobar que el frontend opere en navegadores modernos y se adapte a dispositivos móviles y de escritorio."),
         ("Organizar arquitectura y stack técnico",
          "Mantener backend Python y frontend React en módulos desacoplados, comunicados por API REST JSON, "
          "usando MySQL como motor principal con integridad referencial."),
         ("Implementar respaldo y recuperación",
          "Generar respaldos diarios con RPO máximo de 24 horas y RTO máximo de 4 horas."),
         ("Implementar trazabilidad y observabilidad",
          "Registrar operaciones críticas con actor, fecha, valor previo y valor nuevo cuando aplique, "
          "además de logs, métricas y trazas para diagnosticar fallos operativos y de integración."),
     ]),
]


# =========================================================================
# DOCUMENTO 1 — SRS
# =========================================================================

def doc_srs():
    doc = Document()
    set_doc_lang(doc); set_margins(doc)
    add_cover(doc, entregable="Especificación de Requisitos del Software (SRS)", **COVER_KW)

    # ---- Sección: Propósito ----
    add_h1(doc, "1. Propósito del sistema")
    add_para(doc,
        "Este documento describe la Especificación de Requisitos del Software (SRS) del proyecto "
        "Distrito Urbano — Tienda Digital, una plataforma e-commerce desarrollada como ejercicio "
        "académico de la asignatura Ingeniería de Software bajo metodología Scrum. La aplicación "
        "permite a los clientes navegar un catálogo con identidad visual propia, buscar productos, "
        "configurar variantes (color y talla), agregar al carrito y completar la compra mediante una "
        "pasarela de pagos. A los administradores les ofrece un panel con gestión de catálogo, "
        "inventario con alertas de stock mínimo, pedidos operativos, finanzas con indicadores y "
        "exportación de reportes, reseñas con moderación y configuración general de la tienda."
    )
    add_para(doc,
        "El propósito del SRS es servir como contrato funcional entre el Product Owner, el equipo "
        "Scrum (Daniel Villamizar — frontend, Santiago Pérez — backend, Tomás Urieles — Scrum "
        "Master e infraestructura) y los stakeholders académicos. Documenta los requisitos "
        "funcionales con identificador único, los requisitos no funcionales, las restricciones del "
        "dominio, los actores que interactúan con el sistema y el modelo de casos de uso. "
        "Acompaña al equipo durante el desarrollo y sirve como referencia para la verificación y "
        "validación del producto al cierre de cada sprint."
    )

    # Alcance
    add_h2(doc, "1.1 Alcance")
    add_para(doc,
        "El alcance del MVP cubre los flujos funcionales esenciales de un ecommerce: gestión de "
        "cuentas y sesiones, vitrina pública con identidad visual, búsqueda y ficha de producto con "
        "variantes, carrito y checkout con cálculo de totales, integración con pasarela de pago "
        "(simulada para fines académicos con cuatro escenarios: aprobado, rechazado, pendiente y "
        "fallido), historial y seguimiento de pedidos con notificaciones, panel administrativo de "
        "catálogo e inventario, módulo financiero con COGS y reportes, configuración de la tienda y "
        "reseñas validadas por compra entregada."
    )

    # Definiciones
    add_h2(doc, "1.2 Definiciones, acrónimos y abreviaturas")
    defs = [
        ("SRS", "Software Requirements Specification (Especificación de Requisitos del Software)."),
        ("RF", "Requisito Funcional."),
        ("RNF", "Requisito No Funcional."),
        ("JWT", "JSON Web Token. Mecanismo de autenticación basado en tokens firmados."),
        ("SAGA", "Patrón de transacciones distribuidas con compensaciones. Aquí se aplica de forma orquestada y síncrona en el checkout."),
        ("COGS", "Cost of Goods Sold. Costo de los productos efectivamente vendidos en un período."),
        ("MVP", "Minimum Viable Product. Versión mínima funcional del producto."),
        ("SKU", "Stock Keeping Unit. Identificador único por variante de inventario."),
        ("RPO / RTO", "Recovery Point Objective y Recovery Time Objective. Métricas de continuidad y recuperación de datos."),
    ]
    for term, expl in defs:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.6)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(f"{term}: "); set_run(run, bold=True, size=10.5, color=BRAND_DARK)
        run = p.add_run(expl); set_run(run, size=10.5)

    # ---- Sección: Visión general ----
    add_h1(doc, "2. Visión general del producto")
    add_para(doc,
        "Distrito Urbano es una marca ficticia de moda urbana usada como dominio del proyecto. "
        "La plataforma se construye con frontend React + Vite, backend distribuido en 5 microservicios "
        "FastAPI (Auth, Catalog, Inventory, Commerce y Payment), persistencia MySQL con esquemas "
        "aislados por servicio, Redis para caché y locks distribuidos, Nginx como API Gateway y "
        "Mailhog como servidor SMTP local. Toda la solución corre localmente sobre Docker Compose."
    )

    # ---- Sección: Restricciones del dominio ----
    add_h1(doc, "3. Restricciones del dominio")
    add_para(doc, "El sistema debe cumplir las siguientes restricciones derivadas del dominio comercial y técnico:")

    restrs = [
        ("Stack tecnológico fijo",
         "Backend en Python 3.11+ (FastAPI), frontend React 18 + Vite, base de datos MySQL 8.4, "
         "comunicación REST con JSON, autenticación JWT HS256 con secreto compartido entre servicios."),
        ("Una sola pasarela de pago",
         "La pasarela es simulada y devuelve únicamente cuatro escenarios: APPROVED, REJECTED, "
         "PENDING y FAILED. No se almacena información de tarjeta de crédito."),
        ("Reglas de negocio sobre pedidos",
         "Solo se persiste un pedido (Order) cuando el checkout termina en pago APPROVED. Los "
         "intentos fallidos quedan en una tabla separada (FailedCheckoutAttempt) para auditoría sin "
         "contaminar el panel administrativo."),
        ("Reseñas restringidas",
         "Únicamente se permiten reseñas sobre productos comprados y entregados. Toda reseña entra "
         "en estado pendiente y debe ser aprobada por el administrador antes de publicarse."),
        ("Inventario con SKU único",
         "Cada variante de producto tiene un SKU único global. La combinación (producto, color, talla) "
         "es también única para impedir duplicados visibles al cliente."),
        ("Idioma y moneda",
         "Interfaz en español de Colombia (es-CO). Moneda única: peso colombiano (COP). Zona horaria "
         "America/Bogota."),
        ("Roles cerrados",
         "Únicamente existen dos roles: cliente (customer) y administrador (admin). No hay roles intermedios."),
        ("Carrito por usuario",
         "Cada cliente autenticado posee un único carrito en estado abierto. No se admiten carritos anónimos."),
        ("Trazabilidad",
         "Toda acción crítica (login, registro, checkout, transición de pedido, aprobación de reseña) "
         "se registra con actor, fecha y correlation-id."),
    ]
    for t, d in restrs:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(f"{t}. "); set_run(run, bold=True, size=10.5, color=BRAND_DARK)
        run = p.add_run(d); set_run(run, size=10.5)

    # ---- Sección: Actores del sistema ----
    add_h1(doc, "4. Actores del sistema")
    add_para(doc, "Se identifican los siguientes actores que interactúan con la plataforma:")

    t = add_table_header(doc, ["Actor", "Tipo", "Responsabilidades principales"], widths_cm=[4, 2.5, 9])
    actores = [
        ("Cliente", "Humano",
         "Registrarse, iniciar sesión, explorar catálogo, buscar y consultar productos, gestionar "
         "carrito, realizar checkout, consultar pedidos, reseñar productos entregados y recibir notificaciones."),
        ("Administrador", "Humano",
         "Gestionar catálogo (categorías, productos, imágenes), administrar inventario (variantes, "
         "movimientos, alertas), gestionar pedidos operativos, aprobar reseñas, consultar finanzas "
         "y reportes, configurar la tienda y consultar la bitácora de auditoría."),
        ("Pasarela de Pago", "Sistema externo",
         "Recibir solicitudes de cobro emitidas por el servicio de pagos y devolver la autorización, "
         "rechazo, pendiente o fallo de la transacción. Implementada como simulador para el MVP."),
        ("Servicio de Correo (SMTP)", "Sistema externo",
         "Recibir y entregar correos transaccionales emitidos por el sistema: bienvenida al "
         "registrarse, confirmación de pedido, cambio de estado y notificaciones administrativas."),
    ]
    for fila in actores:
        add_table_row(t, fila)

    # ---- Sección: Requisitos funcionales ----
    add_h1(doc, "5. Requisitos funcionales")
    add_para(doc,
        "Los requisitos funcionales se identifican con la convención RF-NN. Cada requisito presenta "
        "su título, criterio de aceptación en formato de historia de usuario y las tareas (tasks) "
        "asociadas que materializan su implementación.")

    for rf_id, titulo, decl, ac, tasks in RF:
        add_h2(doc, f"{rf_id} — {titulo}")
        # Tabla resumen
        t = add_table_header(doc, ["Campo", "Detalle"], widths_cm=[3.5, 13])
        add_table_row(t, ["ID", rf_id])
        add_table_row(t, ["Declaración", decl])
        add_table_row(t, ["Criterio de aceptación", ac])
        # Tabla de tareas
        add_h3(doc, "Tareas asociadas")
        tt = add_table_header(doc, ["#", "Tarea", "Descripción"], widths_cm=[1, 5, 11])
        for i, (tt_name, tt_desc) in enumerate(tasks, 1):
            add_table_row(tt, [str(i), tt_name, tt_desc])

    # ---- Sección: Requisitos no funcionales ----
    add_h1(doc, "6. Requisitos no funcionales")
    add_para(doc,
        "Los requisitos no funcionales describen las propiedades de calidad de la solución. "
        "Se identifican con la convención RNF-NN y abarcan rendimiento, seguridad y aspectos "
        "transversales como usabilidad, respaldo y trazabilidad.")

    for rnf_id, titulo, decl, ac, tasks in RNF:
        add_h2(doc, f"{rnf_id} — {titulo}")
        t = add_table_header(doc, ["Campo", "Detalle"], widths_cm=[3.5, 13])
        add_table_row(t, ["ID", rnf_id])
        add_table_row(t, ["Declaración", decl])
        add_table_row(t, ["Criterio de aceptación", ac])
        add_h3(doc, "Tareas asociadas")
        tt = add_table_header(doc, ["#", "Tarea", "Descripción"], widths_cm=[1, 5, 11])
        for i, (tt_name, tt_desc) in enumerate(tasks, 1):
            add_table_row(tt, [str(i), tt_name, tt_desc])

    # ---- Sección: Visión scrum / metodología ----
    add_h1(doc, "7. Contexto Scrum del proyecto")
    add_para(doc,
        "El proyecto se desarrolla siguiendo el marco Scrum para simular un entorno de trabajo "
        "ágil. El equipo está conformado por tres integrantes con responsabilidades diferenciadas y "
        "se organiza en sprints semanales con artefactos formales del marco. Cada requisito "
        "funcional se descompone en tareas (tasks) trazables a commits en el repositorio.")

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    run = p.add_run("Roles del equipo:"); set_run(run, bold=True, size=11, color=BRAND_DARK)

    add_bullet(doc, "Product Owner (rol simulado): define los requisitos, prioriza el backlog y valida los entregables.")
    add_bullet(doc, "Scrum Master — Tomás Urieles: facilita el proceso, gestiona la integración, "
                    "infraestructura, documentación, pruebas y los requisitos no funcionales.")
    add_bullet(doc, "Equipo de desarrollo: Daniel Villamizar (frontend React, experiencia de usuario, "
                    "componentes) y Santiago Pérez (backend FastAPI, base de datos, APIs, lógica de "
                    "negocio, integración con pasarela y servicios de correo).")

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    run = p.add_run("Artefactos:"); set_run(run, bold=True, size=11, color=BRAND_DARK)
    add_bullet(doc, "Product Backlog: lista priorizada de requisitos funcionales (RF-01 a RF-09).")
    add_bullet(doc, "Sprint Backlog: tareas planificadas para cada sprint, tomadas del backlog.")
    add_bullet(doc, "Incremento: software desplegable al final de cada sprint, integrado en Docker Compose.")
    add_bullet(doc, "Bitácora de auditoría: trazabilidad de acciones críticas (RNF-03).")

    doc.save(BASE / "1_SRS.docx")
    print("  OK  1_SRS.docx")


# =========================================================================
# DOCUMENTO 2 — Diagrama de Casos de Uso
# =========================================================================

def doc_diagrama_casos_uso():
    doc = Document()
    set_doc_lang(doc); set_margins(doc)
    add_cover(doc, entregable="Diagrama de Casos de Uso", **COVER_KW)

    add_h1(doc, "1. Introducción")
    add_para(doc,
        "Este documento presenta el diagrama UML de casos de uso del sistema Distrito Urbano. "
        "El diagrama muestra los actores principales (cliente, administrador y dos sistemas "
        "externos: la pasarela de pago y el servicio SMTP) junto con los casos de uso del sistema "
        "en infinitivo. Se modelan las relaciones <<include>> para los pasos obligatorios y "
        "<<extend>> para los caminos opcionales del flujo.")

    add_h1(doc, "2. Convenciones del diagrama")
    add_bullet(doc, "Las elipses representan casos de uso, nombrados en infinitivo.")
    add_bullet(doc, "Los “stick figures” representan actores. El cliente y el administrador son "
                    "actores humanos; la pasarela de pago y el servicio SMTP son sistemas externos.")
    add_bullet(doc, "Las líneas continuas representan asociaciones actor → caso de uso.")
    add_bullet(doc, "Las flechas punteadas azules indican relación <<include>>: el caso destino se "
                    "ejecuta obligatoriamente como parte del caso origen.")
    add_bullet(doc, "Las flechas punteadas rojas indican relación <<extend>>: el caso destino "
                    "extiende opcionalmente el comportamiento del caso origen en ciertas condiciones.")

    add_h1(doc, "3. Diagrama")
    add_image(doc, IMG / "01_casos_de_uso.png", width_cm=16,
              caption="Diagrama de casos de uso del sistema Distrito Urbano.")

    add_h1(doc, "4. Relaciones <<include>> y <<extend>>")
    add_para(doc, "Se destacan las siguientes relaciones modeladas en el diagrama:")

    t = add_table_header(doc, ["Caso origen", "Relación", "Caso destino", "Significado"],
                          widths_cm=[4, 2.4, 4, 6])
    rels = [
        ("Realizar checkout", "<<include>>", "Gestionar carrito",
         "El checkout requiere haber preparado el carrito previamente."),
        ("Realizar checkout", "<<include>>", "Procesar pago",
         "Todo checkout obliga a invocar la pasarela de pago como paso central."),
        ("Reseñar producto", "<<include>>", "Ver mis pedidos",
         "Para reseñar, el cliente debe acceder a un pedido propio entregado."),
        ("Explorar catálogo", "<<extend>>", "Buscar producto",
         "La búsqueda con texto/filtros extiende la exploración general del catálogo."),
        ("Procesar pago", "<<extend>>", "Recibir notificación",
         "Cuando el pago se confirma, se dispara una notificación al cliente (correo + push)."),
    ]
    for r in rels:
        add_table_row(t, r)

    add_h1(doc, "5. Lista de casos de uso del sistema")
    p = doc.add_paragraph()
    run = p.add_run("Casos de uso del Cliente:"); set_run(run, bold=True, size=11, color=BRAND_DARK)
    cu_cliente = [
        "Registrarse",
        "Iniciar sesión",
        "Explorar catálogo",
        "Buscar producto  (extiende explorar)",
        "Consultar detalle de producto",
        "Gestionar carrito",
        "Realizar checkout  (incluye gestionar carrito y procesar pago)",
        "Procesar pago  (interactúa con la pasarela externa)",
        "Ver mis pedidos",
        "Reseñar producto  (requiere un pedido entregado)",
        "Recibir notificación  (extiende procesar pago)",
    ]
    for cu in cu_cliente:
        add_bullet(doc, cu)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run("Casos de uso del Administrador:"); set_run(run, bold=True, size=11, color=BRAND_DARK)
    cu_admin = [
        "Gestionar catálogo e inventario",
        "Gestionar pedidos operativos (transicionar estados)",
        "Aprobar reseñas",
        "Configurar tienda",
        "Consultar finanzas y reportes",
        "Consultar bitácora de auditoría",
        "Gestionar empleados y gastos",
    ]
    for cu in cu_admin:
        add_bullet(doc, cu)

    doc.save(BASE / "2_Diagrama_Casos_de_Uso.docx")
    print("  OK  2_Diagrama_Casos_de_Uso.docx")


# =========================================================================
# DOCUMENTO 3 — Especificación de Casos de Uso
# =========================================================================

CASOS_DE_USO = [
    {
        "id": "CU-01",
        "nombre": "Registrarse e iniciar sesión",
        "actor": "Cliente",
        "rf": "RF-01",
        "objetivo": "Permitir que un visitante cree una cuenta en la tienda digital y obtenga "
                    "credenciales para acceder a las funcionalidades reservadas a usuarios autenticados "
                    "(carrito, checkout, historial, reseñas).",
        "precondicion":
            "El visitante posee una dirección de correo electrónico válida y no registrada "
            "previamente en la plataforma. El sistema está disponible (servicio de autenticación "
            "saludable).",
        "postcondicion":
            "Se crea un nuevo usuario con rol cliente en la base de datos auth_db, se emite un "
            "token JWT de acceso (access_token) y otro de refresco (refresh_token), se registra "
            "el evento en la bitácora de accesos y se envía un correo de bienvenida.",
        "flujo": [
            "El visitante navega a la ruta /registro de la aplicación.",
            "El sistema presenta el formulario con los campos: nombre, correo, teléfono y contraseña.",
            "El visitante diligencia el formulario y presiona “Registrarme”.",
            "El frontend envía POST /api/auth/register al API Gateway con el cuerpo en JSON.",
            "El Gateway enruta la solicitud al Auth Service.",
            "El Auth Service valida el formato de los campos: contraseña fuerte (≥8 caracteres, "
            "mayúsculas, minúsculas, dígito, símbolo), correo único, teléfono numérico.",
            "El servicio almacena el usuario con contraseña hasheada (bcrypt) y rol = customer.",
            "El servicio genera el par de tokens JWT firmados con HS256.",
            "El servicio registra el evento “register” en la tabla access_logs con IP, user_agent "
            "y correlation_id.",
            "El servicio envía un correo de bienvenida al servidor SMTP (Mailhog).",
            "El servicio devuelve 201 Created con los tokens y los datos públicos del usuario.",
            "El frontend persiste el access_token en almacenamiento local y redirige al catálogo.",
        ],
        "alternativos": [
            {
                "id": "FA-01",
                "titulo": "Correo electrónico ya registrado",
                "pasos": [
                    "En el paso 6, si el correo ya existe en la base de datos, el servicio devuelve "
                    "409 Conflict con el detalle “Email ya registrado”.",
                    "El frontend muestra al visitante el mensaje “Ya existe una cuenta con ese "
                    "correo”. Sugerencia: iniciar sesión.",
                    "El caso de uso termina sin crear cuenta.",
                ],
            },
            {
                "id": "FA-02",
                "titulo": "Contraseña débil",
                "pasos": [
                    "En el paso 6, si la contraseña no cumple los criterios de fortaleza, el servicio "
                    "devuelve 422 Unprocessable Entity con la lista de criterios incumplidos.",
                    "El frontend resalta los criterios faltantes (rojo) y permite al visitante corregir.",
                    "El caso de uso retorna al paso 3.",
                ],
            },
            {
                "id": "FA-03",
                "titulo": "Demasiados intentos desde la misma IP (rate limit)",
                "pasos": [
                    "Si el visitante intenta registrarse más de 5 veces por minuto desde la misma IP, "
                    "el API Gateway responde 429 Too Many Requests sin tocar el Auth Service.",
                    "El frontend muestra “Demasiados intentos. Intenta de nuevo en un minuto”.",
                ],
            },
        ],
    },
    {
        "id": "CU-02",
        "nombre": "Buscar y consultar producto",
        "actor": "Cliente (autenticado o no)",
        "rf": "RF-02, RF-03",
        "objetivo": "Permitir a cualquier visitante encontrar productos del catálogo, filtrarlos por "
                    "categoría, precio o disponibilidad, y consultar la ficha completa con sus variantes.",
        "precondicion":
            "El servicio Catalog está disponible y hay productos publicados (campo published=true, "
            "archived=false). El cliente NO necesita estar autenticado para este caso.",
        "postcondicion":
            "El cliente visualiza la lista filtrada de productos y/o el detalle del producto "
            "seleccionado con sus variantes (color, talla) y disponibilidad real en tiempo real.",
        "flujo": [
            "El visitante accede a /catalogo desde el navbar de la tienda.",
            "El frontend hace GET /api/products?... al gateway, sin parámetros.",
            "Catalog Service consulta su esquema catalog_db y obtiene los productos publicados.",
            "Catalog Service consulta a Inventory Service GET /stock-summary para enriquecer cada "
            "producto con su stock agregado.",
            "Catalog Service responde con la lista, incluyendo stock, variant_count y un flag "
            "inventory_available.",
            "El frontend renderiza las tarjetas del catálogo, marcando como “AGOTADO” los productos "
            "sin stock (con un overlay grisáceo).",
            "El visitante escribe texto en el buscador y/o selecciona filtros (categoría, precio, “Con stock”).",
            "El frontend hace una nueva solicitud GET /api/products?q=…&category_id=…&min_price=…",
            "El visitante hace clic en una tarjeta para abrir el detalle.",
            "El frontend solicita GET /api/products/{id} y GET /api/reviews/product/{id}.",
            "Catalog Service obtiene el detalle del producto y llama a Inventory Service "
            "GET /products/{id}/variants para obtener variantes con stock real, costo, precio, "
            "color y color_hex.",
            "El frontend agrupa las variantes por color, mostrando un selector tipo Nike: primero el "
            "color (bolita visual) y luego las tallas disponibles dentro de cada color.",
            "El visitante examina las reseñas aprobadas de otros clientes (sello “compra verificada”).",
        ],
        "alternativos": [
            {
                "id": "FA-01",
                "titulo": "Producto sin variantes registradas",
                "pasos": [
                    "En el paso 11, si Inventory Service devuelve lista vacía o variant_count = 0, "
                    "el frontend muestra el producto como AGOTADO (efecto grisáceo) y deshabilita "
                    "el botón “Agregar al carrito”.",
                ],
            },
            {
                "id": "FA-02",
                "titulo": "Inventory Service temporalmente caído",
                "pasos": [
                    "Si la llamada a Inventory falla, Catalog Service responde la lista de productos "
                    "con inventory_available = false en cada item.",
                    "El frontend muestra el mensaje “Consultar” en lugar del stock y permite al "
                    "visitante seguir navegando.",
                    "El caso de uso prosigue en modo degradado: el detalle del producto muestra el "
                    "precio base pero sin variantes confirmadas.",
                ],
            },
        ],
    },
    {
        "id": "CU-03",
        "nombre": "Realizar compra (carrito + checkout + pago)",
        "actor": "Cliente",
        "rf": "RF-04, RF-05, RF-06",
        "objetivo": "Permitir que un cliente autenticado seleccione productos, los agregue al carrito, "
                    "ingrese los datos de entrega y finalice la compra mediante una pasarela de pago, "
                    "creando un pedido con estado PAID.",
        "precondicion":
            "El cliente está autenticado (token JWT válido), seleccionó al menos una variante con "
            "stock disponible y tiene saldo/método de pago válido en la pasarela mock.",
        "postcondicion":
            "Si el pago es aprobado: se crea una Order en commerce_db con estado PAID, se descuenta "
            "el stock real en Inventory, se guarda el unit_cost (snapshot) en cada OrderItem para "
            "cálculo de COGS, se envía notificación in-app y correo al cliente.\n"
            "Si el pago falla: NO se crea Order. Se libera la reserva de stock y se registra el "
            "intento en FailedCheckoutAttempt.",
        "flujo": [
            "El cliente abre la página del producto y selecciona color y talla.",
            "El cliente hace clic en “Agregar al carrito” con la cantidad deseada.",
            "El frontend envía POST /api/cart/items {variant_id, quantity} a Commerce Service.",
            "Commerce Service valida el stock disponible llamando a Inventory y persiste el ítem en el carrito del usuario.",
            "El cliente abre /carrito y revisa los ítems, totales y opcionalmente ajusta cantidades.",
            "El cliente pulsa “Continuar al checkout”.",
            "El frontend lo lleva al formulario de datos de entrega (nombre, dirección, ciudad, "
            "teléfono, email, documento de facturación).",
            "El cliente diligencia el formulario y confirma el pago.",
            "El frontend envía POST /api/checkout con los datos de entrega + header Idempotency-Key.",
            "Commerce Service inicia la SAGA orquestada síncrona: paso 1 — POST /reserve a Inventory "
            "Service. Inventory toma lock distribuido Redis por variante, hace SELECT FOR UPDATE en MySQL "
            "y reserva el stock por 15 minutos.",
            "Paso 2 — Commerce llama POST /payments a Payment Service. Payment verifica que el Circuit "
            "Breaker esté CLOSED y emite POST /charge a la pasarela mock.",
            "La pasarela responde APPROVED. Payment devuelve {status: APPROVED, ref: AUTH-xxx}.",
            "Paso 3 — Commerce llama POST /confirm/{order_id} a Inventory para descontar el stock "
            "definitivamente (stock -= qty, reserved_stock -= qty).",
            "Commerce persiste la Order con status=PAID, OrderItems con unit_cost del snapshot de Inventory, "
            "OrderStatusHistory y OrderAuditLog.",
            "Commerce inserta una Notification para el cliente y envía correo vía SMTP.",
            "Commerce marca el carrito como “checked_out” (vacía el carrito abierto del usuario).",
            "Commerce devuelve 201 Created con {order_id, order_code, status: PAID, total}.",
            "El frontend redirige a la pantalla “¡Pago aprobado!” con el código del pedido y botones "
            "para ver detalle o seguir comprando.",
        ],
        "alternativos": [
            {
                "id": "FA-01",
                "titulo": "Stock insuficiente al reservar",
                "pasos": [
                    "En el paso 10, si Inventory devuelve 409 Conflict con la lista de variantes sin stock, "
                    "Commerce registra el intento en FailedCheckoutAttempt(reason_code='out_of_stock'), "
                    "NO crea Order y devuelve HTTP 409 con el cuerpo {code: 'out_of_stock', unavailable: [...]}.",
                    "El frontend muestra la pantalla de error con el detalle de cada variante sin stock y "
                    "ofrece “Intentar de nuevo” o “Volver al carrito”.",
                ],
            },
            {
                "id": "FA-02",
                "titulo": "Pago rechazado por la pasarela",
                "pasos": [
                    "En el paso 12, si la pasarela devuelve REJECTED, Commerce lanza la compensación: "
                    "POST /release a Inventory para liberar la reserva.",
                    "Commerce registra el intento en FailedCheckoutAttempt(reason_code='payment_rejected'), "
                    "envía notificación + correo al cliente y devuelve HTTP 402 con {code: 'payment_rejected'}.",
                    "El frontend muestra la pantalla “Pago rechazado” con la razón devuelta por la "
                    "pasarela y el botón “Intentar con otro método”.",
                ],
            },
            {
                "id": "FA-03",
                "titulo": "Pasarela caída — Circuit Breaker abierto",
                "pasos": [
                    "En el paso 11, si la pasarela ha fallado 5 veces consecutivas en los últimos 60 segundos, "
                    "el Circuit Breaker de Payment Service se encuentra en estado OPEN.",
                    "Payment responde inmediatamente 503 Service Unavailable sin tocar la pasarela.",
                    "Commerce libera la reserva, registra el intento en FailedCheckoutAttempt y devuelve "
                    "503 con {code: 'payment_unavailable'}.",
                    "El frontend muestra “La pasarela está temporalmente fuera de servicio” con botón "
                    "“Intentar de nuevo en unos minutos”.",
                ],
            },
            {
                "id": "FA-04",
                "titulo": "Reintento con la misma Idempotency-Key",
                "pasos": [
                    "Si la solicitud POST /checkout llega dos veces con la misma Idempotency-Key del "
                    "mismo usuario (por ejemplo, doble clic), Commerce detecta el correlation_id ya "
                    "procesado y devuelve la Order ya creada en vez de procesarla de nuevo.",
                    "No hay doble cobro ni doble descuento de stock.",
                ],
            },
        ],
    },
    {
        "id": "CU-04",
        "nombre": "Gestionar pedido como administrador",
        "actor": "Administrador",
        "rf": "RF-06, RF-07",
        "objetivo": "Permitir al administrador consultar los pedidos pagados y avanzar su estado "
                    "operativo (preparación → envío → entrega) o cancelarlos si es necesario, "
                    "notificando al cliente automáticamente en cada transición.",
        "precondicion":
            "El administrador está autenticado con un JWT que contiene el rol admin. Existe al "
            "menos un pedido en estado PAID o posterior en el sistema.",
        "postcondicion":
            "El estado del pedido es actualizado, se inserta un registro en OrderStatusHistory con "
            "el actor (admin_id), la transición de estado y el timestamp. Se inserta un evento en "
            "OrderAuditLog y se notifica al cliente vía notificación in-app y correo.",
        "flujo": [
            "El administrador inicia sesión en /login con sus credenciales de admin.",
            "El sistema lo redirige a /admin (panel administrativo).",
            "El administrador hace clic en “Pedidos” en el menú lateral.",
            "El frontend llama a GET /api/admin/orders. Commerce Service devuelve solo los pedidos "
            "en estados operativos (PAID, EN_PREPARACION, ENVIADO, ENTREGADO, CANCELADA).",
            "El frontend muestra el listado con filtros por estado, búsqueda por código y datos de cliente.",
            "El administrador hace clic en un pedido en estado PAID.",
            "El frontend abre el detalle del pedido en un modal.",
            "El administrador selecciona “En preparación” en el dropdown de estado.",
            "El frontend envía PATCH /api/admin/orders/{id}/status {new_status: 'EN_PREPARACION'}.",
            "Commerce Service valida que la transición sea legal (PAID → EN_PREPARACION está permitido) "
            "y actualiza el estado.",
            "El servicio inserta una entrada en OrderStatusHistory con from_status, to_status, "
            "changed_by=admin_id y notes.",
            "Inserta un evento en OrderAuditLog con action=status_change_PAID_to_EN_PREPARACION.",
            "Crea una Notification para el cliente con título “Tu pedido está en preparación”.",
            "Envía un correo SMTP al cliente con el cambio.",
            "Devuelve 200 OK con el pedido actualizado.",
            "El frontend refresca el listado y muestra un toast de confirmación.",
            "El administrador puede repetir los pasos 6-16 para EN_PREPARACION → ENVIADO → ENTREGADO.",
        ],
        "alternativos": [
            {
                "id": "FA-01",
                "titulo": "Transición de estado inválida",
                "pasos": [
                    "Si el administrador intenta retroceder un estado (por ejemplo, ENVIADO → PAID) o "
                    "saltar etapas (PAID → ENVIADO), Commerce devuelve 409 Conflict con la lista de "
                    "transiciones legales desde el estado actual.",
                    "El frontend muestra un toast de error y mantiene el estado anterior.",
                ],
            },
            {
                "id": "FA-02",
                "titulo": "Cancelación de pedido",
                "pasos": [
                    "En cualquier momento del flujo, el administrador puede transicionar el pedido a "
                    "CANCELADA desde PAID o EN_PREPARACION (no se permite cancelar pedidos ya enviados).",
                    "La cancelación dispara la misma cadena: status_history + audit_log + notification + correo.",
                ],
            },
        ],
    },
    {
        "id": "CU-05",
        "nombre": "Reseñar producto comprado",
        "actor": "Cliente",
        "rf": "RF-09",
        "objetivo": "Permitir al cliente dejar una reseña con calificación y comentario sobre un "
                    "producto que ya recibió. La reseña pasa por un proceso de aprobación administrativa "
                    "antes de hacerse pública.",
        "precondicion":
            "El cliente está autenticado y posee al menos un pedido en estado ENTREGADO. El "
            "producto que desea reseñar pertenece a ese pedido y no tiene una reseña previa del "
            "mismo cliente para el mismo pedido.",
        "postcondicion":
            "Se crea una Review con approved=false en commerce_db. La reseña queda pendiente "
            "hasta la aprobación del administrador. Una vez aprobada, se publica en la ficha del "
            "producto y se recalcula el rating promedio en el RatingSummary de Catalog.",
        "flujo": [
            "El cliente abre /mis-pedidos y selecciona un pedido en estado ENTREGADO.",
            "En el detalle del pedido aparece un botón “Dejar reseña” por cada producto comprado.",
            "El cliente hace clic en “Dejar reseña” del producto que desea reseñar.",
            "El frontend lo lleva a /resenas/{orderId}/{productId}.",
            "El cliente selecciona una calificación (1 a 5 estrellas) con el star picker.",
            "El cliente escribe un comentario en el textarea (con contador de caracteres).",
            "El cliente hace clic en “Publicar reseña”.",
            "El frontend envía POST /api/reviews {order_id, product_id, rating, comment}.",
            "Commerce Service verifica que (order_id, user_id) corresponde a un pedido ENTREGADO del "
            "cliente y que el product_id existe entre los OrderItems.",
            "Verifica que no exista ya una Review para (user_id, product_id, order_id).",
            "Persiste la Review con approved=false.",
            "Devuelve 201 Created con la reseña creada.",
            "El frontend muestra un toast “Reseña publicada. Quedará pendiente de aprobación”.",
            "El frontend vuelve al detalle del pedido y marca el producto como “Reseñado” (badge verde).",
            "Asincrónicamente, el administrador entra a /admin/resenas y ve la reseña como pendiente.",
            "El administrador la aprueba (PATCH /api/admin/reviews/{id}/approve).",
            "Commerce recalcula el rating promedio y llama a PUT /admin/products/{id}/rating en "
            "Catalog para actualizar el RatingSummary.",
            "La reseña aparece públicamente en la ficha del producto con el sello “Compra verificada”.",
        ],
        "alternativos": [
            {
                "id": "FA-01",
                "titulo": "Producto no comprado / no entregado",
                "pasos": [
                    "Si el cliente intenta enviar una reseña de un producto que no compró o cuyo "
                    "pedido no está en estado ENTREGADO, el servicio responde 409 Conflict con el "
                    "mensaje “Solo puedes reseñar productos comprados y entregados”.",
                ],
            },
            {
                "id": "FA-02",
                "titulo": "Reseña duplicada",
                "pasos": [
                    "Si el cliente ya envió una reseña para ese (producto, pedido), el servicio "
                    "devuelve 409 Conflict con el mensaje “Ya reseñaste este producto para este pedido”.",
                ],
            },
            {
                "id": "FA-03",
                "titulo": "El administrador rechaza la reseña",
                "pasos": [
                    "Si el contenido de la reseña no cumple las políticas de la tienda, el "
                    "administrador la elimina (DELETE /api/admin/reviews/{id}).",
                    "Si la reseña ya estaba aprobada, Commerce recalcula el rating promedio descontándola.",
                    "Si estaba pendiente, simplemente se borra sin afectar el rating.",
                ],
            },
        ],
    },
    {
        "id": "CU-06",
        "nombre": "Consultar finanzas y exportar reportes",
        "actor": "Administrador",
        "rf": "RF-08",
        "objetivo": "Permitir al administrador consultar los indicadores financieros del negocio en "
                    "un período seleccionado, visualizar gráficas y exportar un reporte PDF/CSV para "
                    "presentación o archivo contable.",
        "precondicion":
            "El administrador está autenticado con rol admin. Existen pedidos con estado PAID o "
            "posterior en el período consultado. El backend ha capturado el unit_cost en cada "
            "OrderItem al momento del checkout (necesario para calcular COGS y margen real).",
        "postcondicion":
            "El administrador visualiza KPIs, gráficas y tabla detalle por período. Si exporta, se "
            "genera un archivo CSV descargable o una ventana imprimible con formato PDF.",
        "flujo": [
            "El administrador inicia sesión y entra al menú “Finanzas”.",
            "El frontend solicita GET /api/admin/finance/summary?granularity=month al gateway.",
            "Commerce Service consulta la base commerce_db para calcular:\n"
            "  • ventas brutas = SUM(Order.total) WHERE status IN ('PAID', 'EN_PREPARACION', 'ENVIADO', 'ENTREGADO')\n"
            "  • COGS = SUM(OrderItem.unit_cost * quantity)\n"
            "  • margen bruto = ventas - COGS\n"
            "  • gastos operativos = SUM(Expense.amount)\n"
            "  • nómina = SUM(Employee.salary) WHERE employment_status='active'\n"
            "  • utilidad neta = margen bruto - gastos - nómina\n"
            "  • timeseries: misma agregación pero agrupada por día/mes/año.",
            "Responde con un JSON estructurado.",
            "El frontend renderiza 4 KPIs (ventas, COGS, margen bruto, utilidad neta), una gráfica de "
            "líneas con la evolución del período, una gráfica de torta con la distribución financiera "
            "(COGS / gastos operativos / nómina / utilidad), una gráfica de barras con productos más "
            "vendidos y dos tablas resumen.",
            "El administrador cambia el filtro de período (Hoy / 7 días / Mes / Trimestre / Año / Todo) "
            "o ajusta las fechas manualmente.",
            "El frontend hace una nueva llamada con los parámetros actualizados y vuelve a renderizar.",
            "El administrador hace clic en “Reporte PDF”.",
            "El frontend genera dinámicamente una ventana HTML imprimible con: hero corporativo, KPIs, "
            "estado de rentabilidad, evolución del período en tabla, productos más vendidos y gastos por tipo.",
            "El administrador imprime o guarda como PDF desde el cuadro de diálogo del navegador.",
        ],
        "alternativos": [
            {
                "id": "FA-01",
                "titulo": "Exportar CSV en vez de PDF",
                "pasos": [
                    "Si el administrador hace clic en “CSV” en lugar de “Reporte PDF”, el frontend "
                    "genera un Blob con BOM UTF-8 (para Excel) que contiene todos los KPIs, la "
                    "timeseries y las tablas en formato delimitado por punto y coma.",
                ],
            },
            {
                "id": "FA-02",
                "titulo": "Utilidad neta negativa",
                "pasos": [
                    "Si la utilidad neta es < 0 en el período seleccionado, el frontend muestra una "
                    "alerta roja en la parte superior: “Utilidad neta negativa: los costos superan las "
                    "ventas. Revisa gastos y nómina.”",
                    "El reporte PDF muestra un badge rojo “En pérdidas” en el bloque de resultado operativo.",
                ],
            },
            {
                "id": "FA-03",
                "titulo": "Pedidos históricos sin unit_cost",
                "pasos": [
                    "Si la base contiene pedidos antiguos creados antes de la migración a captura de "
                    "unit_cost, el administrador puede ejecutar POST /api/admin/maintenance/backfill-costs "
                    "para que Commerce recorra esos OrderItems y rellene los costos consultando a Inventory.",
                ],
            },
        ],
    },
]


def doc_especificacion_casos_uso():
    doc = Document()
    set_doc_lang(doc); set_margins(doc)
    add_cover(doc, entregable="Especificación de Casos de Uso", **COVER_KW)

    add_h1(doc, "1. Introducción")
    add_para(doc,
        "Este documento detalla la especificación textual de los casos de uso principales del "
        "sistema Distrito Urbano. Para cada caso de uso se identifica: actor, requisito funcional "
        "asociado, objetivo, precondición, postcondición, flujo principal paso a paso y flujos "
        "alternativos (caminos no-felices).")

    add_h1(doc, "2. Convenciones de la especificación")
    add_bullet(doc, "Identificación: cada caso de uso usa el formato CU-NN.")
    add_bullet(doc, "Precondición: condiciones que deben cumplirse ANTES de iniciar el caso de uso.")
    add_bullet(doc, "Postcondición: estado del sistema DESPUÉS de la ejecución exitosa.")
    add_bullet(doc, "Flujo principal: secuencia numerada del camino feliz.")
    add_bullet(doc, "Flujos alternativos (FA): caminos no-felices o variaciones; cada uno indica desde "
                    "qué paso del flujo principal se desvía.")

    add_h1(doc, "3. Catálogo de casos de uso")
    t = add_table_header(doc, ["ID", "Caso de uso", "Actor principal", "RF asociado"],
                          widths_cm=[1.8, 7, 4, 3.5])
    for cu in CASOS_DE_USO:
        add_table_row(t, [cu["id"], cu["nombre"], cu["actor"], cu["rf"]])

    # Detalle de cada caso de uso
    for cu in CASOS_DE_USO:
        add_h1(doc, f"4. {cu['id']} — {cu['nombre']}")

        # Tabla de identificación
        t = add_table_header(doc, ["Campo", "Detalle"], widths_cm=[3.5, 13])
        add_table_row(t, ["Identificador",  cu["id"]])
        add_table_row(t, ["Nombre",         cu["nombre"]])
        add_table_row(t, ["Actor principal", cu["actor"]])
        add_table_row(t, ["RF asociado",    cu["rf"]])
        add_table_row(t, ["Objetivo",        cu["objetivo"]])
        add_table_row(t, ["Precondición",    cu["precondicion"]])
        add_table_row(t, ["Postcondición",   cu["postcondicion"]])

        # Flujo principal
        add_h3(doc, "Flujo principal")
        for i, paso in enumerate(cu["flujo"], 1):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.6)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(f"{i}. "); set_run(run, bold=True, size=10.5, color=BRAND)
            run = p.add_run(paso); set_run(run, size=10.5)

        # Flujos alternativos
        if cu["alternativos"]:
            add_h3(doc, "Flujos alternativos")
            for fa in cu["alternativos"]:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(4)
                run = p.add_run(f"{fa['id']} — {fa['titulo']}")
                set_run(run, bold=True, size=11, color=ACCENT)
                for i, paso in enumerate(fa["pasos"], 1):
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Cm(1.2)
                    p.paragraph_format.space_after = Pt(2)
                    run = p.add_run(f"{i}. "); set_run(run, bold=True, size=10.5, color=ACCENT)
                    run = p.add_run(paso); set_run(run, size=10.5)

    doc.save(BASE / "3_Especificacion_Casos_de_Uso.docx")
    print("  OK  3_Especificacion_Casos_de_Uso.docx")


# =========================================================================
# DOCUMENTO 4 — Diagramas de Comportamiento UML
# =========================================================================

def doc_comportamiento():
    doc = Document()
    set_doc_lang(doc); set_margins(doc)
    add_cover(doc, entregable="Diagramas de Comportamiento UML", **COVER_KW)

    add_h1(doc, "1. Introducción")
    add_para(doc,
        "Los diagramas de comportamiento describen cómo se comportan los objetos del sistema a "
        "lo largo del tiempo. Para Distrito Urbano se presentan tres diagramas representativos: "
        "un diagrama de secuencia que muestra la interacción detallada de la SAGA orquestada del "
        "checkout, un diagrama de actividad con el proceso end-to-end de compra desde la perspectiva "
        "del cliente, y un diagrama de estado del ciclo de vida del pedido.")

    add_h1(doc, "2. Diagrama de secuencia — Realizar compra (Checkout SAGA)")
    add_para(doc,
        "Este diagrama detalla la interacción entre los componentes durante el flujo de checkout. "
        "Muestra la coordinación de la SAGA orquestada y síncrona: Commerce Service actúa como "
        "orquestador y llama secuencialmente a Inventory Service (para reservar el stock con lock "
        "distribuido), a Payment Service (que aplica el Circuit Breaker antes de invocar la pasarela) "
        "y vuelve a Inventory (para confirmar la baja de stock) si el pago es aprobado. Las cajas "
        "de activación verticales representan el período en que cada componente está ejecutando "
        "trabajo. Si el pago es rechazado, Commerce invoca la compensación POST /release para "
        "liberar el stock reservado.")

    add_image(doc, IMG / "02_secuencia_checkout.png", width_cm=17,
              caption="Diagrama de secuencia del checkout SAGA orquestada síncrona.")

    add_para(doc,
        "El flujo alternativo (recuadro amarillo del diagrama) cubre los tres escenarios de error: "
        "stock insuficiente (Inventory devuelve 409), pago rechazado por la pasarela (Commerce "
        "ejecuta compensación) y Circuit Breaker en estado OPEN (la pasarela está caída y Payment "
        "responde 503 sin invocarla). En los tres casos, no se persiste una Order en commerce_db: "
        "se registra el intento en FailedCheckoutAttempt y se devuelve un error HTTP claro al cliente.",
        italic=True)

    add_h1(doc, "3. Diagrama de actividad — Proceso de compra del cliente")
    add_para(doc,
        "El diagrama de actividad muestra el flujo de trabajo del cliente desde que inicia sesión "
        "hasta que recibe la confirmación de la compra. Incluye dos puntos de decisión: "
        "(a) si la variante seleccionada tiene stock, el cliente sigue al carrito; en caso "
        "contrario, vuelve a explorar el catálogo; "
        "(b) si la SAGA de checkout es exitosa, el sistema crea la Order y notifica; en caso "
        "contrario, muestra un error claro (sin stock / pago rechazado / pasarela caída) y el "
        "cliente puede reintentar.")

    add_image(doc, IMG / "03_actividad_compra.png", width_cm=12,
              caption="Diagrama de actividad del flujo de compra del cliente.")

    add_h1(doc, "4. Diagrama de estado — Ciclo de vida del Pedido")
    add_para(doc,
        "El diagrama de estado representa el ciclo de vida de la entidad Order desde su creación "
        "hasta su estado terminal. Bajo la política MVP, un pedido solo existe en commerce_db si "
        "el checkout termina exitosamente en estado PAID. A partir de ese estado, el administrador "
        "puede transicionarlo a EN_PREPARACION, luego a ENVIADO y finalmente a ENTREGADO. Desde "
        "PAID o EN_PREPARACION es posible cancelar el pedido (estado CANCELADA). Los estados "
        "ENVIADO y ENTREGADO no admiten cancelación, dado que el producto ya salió del control de "
        "la tienda.")

    add_image(doc, IMG / "04_estado_pedido.png", width_cm=16,
              caption="Diagrama de estado del Pedido.")

    add_para(doc,
        "Los intentos de checkout que no llegan a PAID (out_of_stock, payment_rejected, "
        "payment_unavailable, inventory_unavailable, payment_pending/failed) NO crean instancia de "
        "Order. En su lugar, quedan registrados en la tabla FailedCheckoutAttempt para trazabilidad "
        "y soporte, evitando contaminar el panel administrativo y las métricas financieras.",
        italic=True)

    add_h1(doc, "5. Transiciones válidas del Pedido")
    t = add_table_header(doc, ["Estado actual", "Transiciones permitidas", "Disparador"],
                          widths_cm=[3.5, 5, 8])
    transiciones = [
        ("(inicial)", "→ PAID", "Checkout exitoso: stock reservado + pago APPROVED + stock confirmado."),
        ("PAID", "→ EN_PREPARACION, → CANCELADA",
         "Administrador inicia preparación del pedido o lo cancela."),
        ("EN_PREPARACION", "→ ENVIADO, → CANCELADA",
         "Administrador despacha el pedido al transportador o lo cancela."),
        ("ENVIADO", "→ ENTREGADO",
         "Administrador confirma entrega al cliente."),
        ("ENTREGADO", "(terminal)",
         "Estado final. Habilita al cliente para reseñar productos del pedido."),
        ("CANCELADA", "(terminal)",
         "Estado final por cancelación administrativa."),
    ]
    for r in transiciones:
        add_table_row(t, r)

    doc.save(BASE / "4_Diagramas_Comportamiento.docx")
    print("  OK  4_Diagramas_Comportamiento.docx")


# =========================================================================
# DOCUMENTO 5 — Diagramas Estructurales UML
# =========================================================================

def doc_estructurales():
    doc = Document()
    set_doc_lang(doc); set_margins(doc)
    add_cover(doc, entregable="Diagramas Estructurales UML", **COVER_KW)

    add_h1(doc, "1. Introducción")
    add_para(doc,
        "Los diagramas estructurales describen las partes que componen el sistema y cómo se "
        "relacionan entre sí, independientemente del tiempo. Para Distrito Urbano se presentan "
        "tres diagramas estructurales fundamentales: diagrama de clases (modelo de dominio "
        "agrupado por microservicio), diagrama de componentes (arquitectura microservicios) y "
        "diagrama de despliegue (contenedores Docker en el entorno local).")

    add_h1(doc, "2. Diagrama de clases — Modelo de dominio")
    add_para(doc,
        "El diagrama de clases muestra las entidades persistentes principales del sistema agrupadas "
        "por microservicio. Cada microservicio posee su propio esquema MySQL aislado, por lo que "
        "las asociaciones entre entidades de distintos microservicios son referencias lógicas (líneas "
        "punteadas grises) y NO claves foráneas físicas — esto permite escalar y desplegar cada "
        "microservicio de forma independiente sin acoplar la base de datos.")

    add_image(doc, IMG / "05_clases_dominio.png", width_cm=17,
              caption="Diagrama de clases del modelo de dominio (entidades principales por microservicio).")

    add_h2(doc, "2.1 Convenciones del diagrama")
    add_bullet(doc, "Cada caja representa una entidad persistente con sus atributos principales.")
    add_bullet(doc, "El estereotipo <<ref-logica>> en los atributos indica que el ID apunta a una "
                    "entidad de otro microservicio pero NO es una clave foránea física.")
    add_bullet(doc, "Las líneas continuas representan asociaciones con clave foránea real (mismo esquema).")
    add_bullet(doc, "Las líneas punteadas grises representan referencias lógicas entre microservicios.")
    add_bullet(doc, "El estereotipo <<unique>> indica restricciones de unicidad relevantes.")
    add_bullet(doc, "Los campos marcados (snapshot) son copias del estado al momento de la transacción "
                    "(por ejemplo, product_name en OrderItem se guarda en el momento del checkout para "
                    "que el detalle del pedido no cambie aunque después se modifique el producto).")

    add_h1(doc, "3. Diagrama de componentes — Arquitectura de microservicios")
    add_para(doc,
        "Este diagrama presenta los componentes lógicos del sistema y sus interfaces. La aplicación "
        "se compone de un frontend SPA en React, un API Gateway en Nginx que sirve como punto único "
        "de entrada, cinco microservicios FastAPI con responsabilidades bien definidas, infraestructura "
        "compartida (MySQL multi-schema, Redis, Mailhog, payment-mock) y herramientas auxiliares "
        "(phpMyAdmin).")

    add_image(doc, IMG / "06_componentes.png", width_cm=17,
              caption="Diagrama de componentes — arquitectura de microservicios.")

    add_h2(doc, "3.1 Componentes y responsabilidades")
    t = add_table_header(doc, ["Componente", "Tecnología", "Responsabilidad"],
                          widths_cm=[4, 3.5, 9])
    componentes = [
        ("Frontend SPA", "React + Vite",
         "Renderizado de la interfaz, gestión del state local, comunicación con el gateway via Axios."),
        ("API Gateway", "Nginx",
         "Punto único de entrada, routing por path /api/<svc>, CORS, headers de seguridad, "
         "rate limit en /auth/login y correlation-id."),
        ("Auth Service", "FastAPI :8001",
         "Registro, login, refresh y logout. Emite JWT HS256. Gestiona la bitácora de accesos."),
        ("Catalog Service", "FastAPI :8002",
         "Productos, categorías, mensajes, configuración de tienda y rating. Cache-Aside con Redis."),
        ("Inventory Service", "FastAPI :8003",
         "Variantes (SKU, color, talla), stock, reservas con lock distribuido, movimientos y alertas."),
        ("Commerce Service", "FastAPI :8004",
         "Carrito, checkout (orquestador SAGA), pedidos, reseñas, notificaciones, finanzas y empleados."),
        ("Payment Service", "FastAPI :8005",
         "Cobros contra la pasarela mock. Circuit Breaker (Redis), reintentos exponenciales, reconciler."),
        ("Payment Mock", "FastAPI :9000",
         "Pasarela simulada con 4 escenarios: APPROVED, REJECTED, PENDING, FAILED (controlados por el monto)."),
        ("MySQL 8.4", "Base de datos relacional",
         "Persistencia. Esquemas aislados: auth_db, catalog_db, inventory_db, commerce_db, payments_db."),
        ("Redis 7", "Cache + cluster state",
         "Cache-Aside del catálogo, locks distribuidos para Inventory y estado del Circuit Breaker."),
        ("Mailhog", "SMTP local",
         "Servidor SMTP de pruebas para validar correos transaccionales (puerto 8025 para UI web)."),
        ("phpMyAdmin", "UI de inspección",
         "Acceso visual a los esquemas MySQL para debugging y administración manual."),
    ]
    for r in componentes:
        add_table_row(t, r)

    add_h2(doc, "3.2 Patrones arquitecturales aplicados")
    patrones = [
        ("API Gateway (Nginx)", "Routing por path, headers de seguridad, CORS, rate-limit, correlation-id."),
        ("SSO con JWT compartido", "El secreto HS256 se comparte entre los servicios; ningún servicio "
                                    "llama de vuelta a Auth para validar tokens."),
        ("Database per Service", "5 esquemas MySQL + un usuario por servicio con GRANT exclusivo. "
                                  "Aislamiento real validado por el Conformity Monkey."),
        ("Cache-Aside",          "Catalog cachea categorías, productos y settings en Redis con TTL "
                                  "60–300 s. Si Redis cae, se sirve desde MySQL (modo degradado)."),
        ("Lock distribuido",     "Inventory adquiere SET NX EX en Redis por variante + SELECT FOR "
                                  "UPDATE en MySQL antes de reservar. Liberación atómica via script Lua."),
        ("SAGA orquestada síncrona", "Commerce coordina reserve → charge → confirm/release como "
                                      "transacción distribuida sin commit atómico."),
        ("Circuit Breaker",      "Payment Service rechaza requests cuando 5 fallos en 60 s; "
                                  "estado CLOSED/OPEN/HALF_OPEN persistido en Redis."),
        ("Retry con backoff",    "Payment reintenta errores transitorios (5xx, timeout) con 250 ms / "
                                  "500 ms / 1 s. Distingue REJECTED (negocio) de errores de infraestructura."),
        ("Worker reconciler",    "Payment ejecuta una tarea asíncrona cada 5 min que reintenta los "
                                  "pagos en estado PENDING o FAILED."),
        ("Idempotencia",         "Commerce acepta Idempotency-Key en /checkout y evita doble cobro "
                                  "si el cliente reintenta con la misma clave."),
        ("Bitácora con correlation-id", "El gateway inyecta X-Correlation-Id en cada request y los "
                                         "servicios lo propagan a logs y a las tablas de auditoría."),
    ]
    t = add_table_header(doc, ["Patrón", "Aplicación"], widths_cm=[5, 12])
    for r in patrones:
        add_table_row(t, r)

    add_h1(doc, "4. Diagrama de despliegue — Contenedores Docker")
    add_para(doc,
        "El diagrama de despliegue muestra cómo se distribuyen los componentes en el entorno local. "
        "Toda la aplicación corre sobre Docker Compose, organizada en 12 contenedores conectados a "
        "una red bridge llamada tienda_net. Cada contenedor expone su puerto al host para "
        "facilitar el debug, pero las comunicaciones internas entre servicios usan los nombres de "
        "contenedor en la red Docker.")

    add_image(doc, IMG / "07_despliegue.png", width_cm=17,
              caption="Diagrama de despliegue — contenedores Docker.")

    add_h2(doc, "4.1 Mapa de puertos")
    t = add_table_header(doc, ["Servicio", "Puerto host", "URL local"], widths_cm=[5, 3, 8])
    puertos = [
        ("API Gateway (Nginx)", "80",       "http://localhost"),
        ("Auth Service",        "8001",     "http://localhost:8001"),
        ("Catalog Service",     "8002",     "http://localhost:8002"),
        ("Inventory Service",   "8003",     "http://localhost:8003"),
        ("Commerce Service",    "8004",     "http://localhost:8004"),
        ("Payment Service",     "8005",     "http://localhost:8005"),
        ("Payment Mock",        "9000",     "http://localhost:9000"),
        ("MySQL",               "3306",     "mysql://localhost:3306"),
        ("phpMyAdmin",          "8080",     "http://localhost:8080"),
        ("Redis",               "6379",     "redis://localhost:6379"),
        ("Mailhog (SMTP)",      "1025",     "smtp://localhost:1025"),
        ("Mailhog (UI web)",    "8025",     "http://localhost:8025"),
        ("Frontend Vite dev",   "5173",     "http://localhost:5173"),
    ]
    for r in puertos:
        add_table_row(t, r)

    add_h2(doc, "4.2 Healthchecks y observabilidad")
    add_para(doc,
        "Cada microservicio expone GET /health, que verifica la conectividad con sus dependencias "
        "críticas (MySQL y, cuando aplica, Redis o la pasarela mock). El API Gateway re-expone "
        "los healthchecks como /health/<svc> para que el Doctor Monkey y los probes de orquestación "
        "tengan un único punto de consulta. Los healthchecks están configurados en docker-compose.yml "
        "con interval, timeout y retries adecuados; los servicios dependientes esperan a que sus "
        "dependencias queden saludables antes de aceptar tráfico (condition: service_healthy).")

    doc.save(BASE / "5_Diagramas_Estructurales.docx")
    print("  OK  5_Diagramas_Estructurales.docx")


# =========================================================================
# Main
# =========================================================================

if __name__ == "__main__":
    print("Generando documentos .docx...")
    doc_srs()
    doc_diagrama_casos_uso()
    doc_especificacion_casos_uso()
    doc_comportamiento()
    doc_estructurales()
    print(f"\n.docx generados en: {BASE}")
