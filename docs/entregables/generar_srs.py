"""Generador de un único documento SRS para Distrito Urbano - Tienda Digital.

Mantiene el estilo del SRS original (azul corporativo, headings con barras,
tablas con header coloreado) y embebe TODOS los diagramas UML como bloques de
código PlantUML para que el usuario los renderice manualmente en plantuml.com
o con el plugin de su IDE.

Salida: docs/entregables/SRS_Tienda_Digital.docx (un único archivo).
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = Path(__file__).parent

# Paleta del SRS original (azul corporativo)
BLUE_HEAD    = RGBColor(0x1F, 0x49, 0x7D)    # H1/H2 azul oscuro
BLUE_LIGHT   = "DCE6F1"                       # fondo header de tabla (hex sin #)
BLUE_BORDER  = "9CC2E5"
BRAND_GREEN  = RGBColor(0x1F, 0x7A, 0x5C)    # Distrito Urbano
ACCENT       = RGBColor(0xF5, 0x9E, 0x0B)
GRAY         = RGBColor(0x4C, 0x59, 0x60)
GRAY_DARK    = RGBColor(0x33, 0x33, 0x33)
TEXT         = RGBColor(0x17, 0x20, 0x26)
CODE_BG      = "F1F5F9"
CODE_BORDER  = "94A3B8"


# =========================================================================
# Helpers de estilo
# =========================================================================

def set_doc_lang(doc, lang="es-CO"):
    """Define el idioma base del DOCX para corrector y metadatos de Word."""
    styles_element = doc.styles.element
    for rpr_default in styles_element.iter(qn("w:rPrDefault")):
        rpr = rpr_default.find(qn("w:rPr"))
        if rpr is None:
            rpr = OxmlElement("w:rPr")
            rpr_default.insert(0, rpr)
        for lang_el in list(rpr.findall(qn("w:lang"))):
            rpr.remove(lang_el)
        lang_el = OxmlElement("w:lang")
        lang_el.set(qn("w:val"), lang)
        lang_el.set(qn("w:eastAsia"), lang)
        lang_el.set(qn("w:bidi"), "ar-SA")
        rpr.append(lang_el)


def set_margins(doc, top=2.2, bottom=2.2, left=2.5, right=2.5):
    """Aplica margenes uniformes a todas las secciones del documento."""
    for section in doc.sections:
        section.top_margin = Cm(top)
        section.bottom_margin = Cm(bottom)
        section.left_margin = Cm(left)
        section.right_margin = Cm(right)


def set_run(run, *, bold=False, italic=False, size=11, color=TEXT, font="Calibri"):
    """Centraliza tipografia, color y enfasis para cada run de texto."""
    run.font.name = font
    r = run._element
    r.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def add_para(doc, text, *, bold=False, italic=False, size=11, color=TEXT,
             align=WD_ALIGN_PARAGRAPH.JUSTIFY, before=4, after=6, indent=0):
    """Agrega un parrafo normal con espaciado y estilo consistente."""
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    set_run(run, bold=bold, italic=italic, size=size, color=color)
    return p


def add_h1(doc, text):
    """H1 estilo SRS original — azul oscuro con barra horizontal."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    set_run(run, bold=True, size=20, color=BLUE_HEAD)
    return p


def add_h2(doc, text):
    """Agrega un encabezado H2 con la paleta corporativa del SRS."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run(run, bold=True, size=14, color=BLUE_HEAD)
    return p


def add_h3(doc, text):
    """Agrega un encabezado H3 para subsecciones del SRS."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    set_run(run, bold=True, size=12, color=BLUE_HEAD)
    return p


def add_h4(doc, text):
    """Agrega un subtitulo compacto para bloques de detalle."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    set_run(run, bold=True, size=11, color=GRAY_DARK)
    return p


def add_bullet(doc, text, level=0, bold=False):
    """Agrega un bullet con indentacion por nivel y formato uniforme."""
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.6 + level * 0.6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    set_run(run, size=10.5, bold=bold)
    return p


def _shade_cell(cell, fill_hex):
    """Aplica color de fondo OOXML a una celda de tabla."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def add_table(doc, headers, rows, widths_cm=None, header_fill=BLUE_LIGHT,
              header_color=BLUE_HEAD, font_size=10, header_size=10):
    """Crea tabla estilo SRS original con header azul claro y borde sutil."""
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Header
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        if widths_cm:
            cell.width = Cm(widths_cm[i])
        _shade_cell(cell, header_fill)
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_run(run, bold=True, size=header_size, color=header_color)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Rows
    for row_data in rows:
        row = t.add_row()
        for i, txt in enumerate(row_data):
            c = row.cells[i]
            if widths_cm:
                c.width = Cm(widths_cm[i])
            c.text = ""
            p = c.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(str(txt))
            set_run(run, size=font_size)
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    return t


def add_code_block(doc, code, *, title=None, lang="plantuml"):
    """Embebe un bloque de código (PlantUML) en un recuadro gris con fuente
    monoespaciada. El usuario puede copiar y pegar en plantuml.com."""
    # Título sobre el bloque
    if title:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(f"📄 {title}  ({lang})")
        set_run(run, bold=True, italic=True, size=9.5, color=GRAY)

    # Tabla 1x1 que actúa como recuadro
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell = t.rows[0].cells[0]
    _shade_cell(cell, CODE_BG)
    cell.text = ""

    # Una línea por párrafo, fuente monoespaciada
    lines = code.strip("\n").split("\n")
    for i, line in enumerate(lines):
        if i == 0:
            p = cell.paragraphs[0]
        else:
            p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.15
        # Preservar espacios iniciales
        if line == "":
            run = p.add_run(" ")
        else:
            run = p.add_run(line.replace("\t", "    "))
        run.font.name = "Consolas"
        run.font.size = Pt(8.5)
        run.font.color.rgb = TEXT

    # Caption indicando que es renderizable
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = cap.add_run(
        "💡 Copiar el bloque anterior y pegarlo en https://www.plantuml.com/plantuml/uml/ "
        "(o en el plugin PlantUML del IDE) para renderizar el diagrama."
    )
    set_run(run, italic=True, size=8.5, color=GRAY)
    cap.paragraph_format.space_after = Pt(10)


def add_cover(doc):
    """Portada al estilo del SRS original — azul corporativo + tabla de metadata."""
    # Espacio superior
    for _ in range(2):
        doc.add_paragraph()

    # Título principal
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("SRS – Aplicación de tienda digital")
    set_run(run, bold=True, size=26, color=BLUE_HEAD)

    # Subtítulo
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "Especificación de Requisitos de Software para Distrito Urbano:\n"
        "ecommerce con arquitectura de microservicios, catálogo, "
        "inventario, pedidos y análisis financiero"
    )
    set_run(run, italic=True, size=12, color=GRAY)

    # Barra azul separadora
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(14)
    run = p.add_run("▬" * 40)
    set_run(run, size=12, color=BLUE_HEAD)

    # Tabla de metadata (estilo del original)
    meta = [
        ("Código del documento",      "SRS-TD-EMP-002"),
        ("Estado",                    "Línea base de requisitos (revisión mayo 2026)"),
        ("Fecha",                     "Mayo de 2026"),
        ("Propietario del documento", "Patrocinador del proyecto / Dirección del negocio"),
        ("Audiencia",                 "Dirección, analistas, UX/UI, desarrollo, QA, seguridad y operación"),
        ("Ámbito",                    "Aplicación web/móvil y panel administrativo de una sola empresa (Distrito Urbano)"),
        ("Tecnologías definidas",     "Backend Python (FastAPI, arquitectura microservicios); base de datos MySQL multi-schema; "
                                       "frontend React + Vite; cache y locks Redis; gateway Nginx; Docker Compose para entorno local"),
    ]
    t = doc.add_table(rows=0, cols=2)
    t.style = "Light Grid Accent 1"
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for label, value in meta:
        row = t.add_row()
        c1 = row.cells[0]; c2 = row.cells[1]
        c1.width = Cm(5.5); c2.width = Cm(11)
        _shade_cell(c1, BLUE_LIGHT)
        c1.text = ""; c2.text = ""
        r = c1.paragraphs[0].add_run(label); set_run(r, bold=True, size=10, color=BLUE_HEAD)
        r = c2.paragraphs[0].add_run(value); set_run(r, size=10)
        c1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        c2.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Espacio
    for _ in range(3):
        doc.add_paragraph()

    # Equipo
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Equipo del proyecto:")
    set_run(run, italic=True, size=11, color=GRAY)

    integrantes = [
        "Tomás Enrique Urieles – Scrum Master / Infraestructura",
        "Daniel Enrique Villamizar – Frontend Developer",
        "Santiago Pérez Flórez – Backend Developer",
    ]
    for i in integrantes:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(i)
        set_run(run, italic=True, size=11, color=TEXT)

    # Cliente
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Empresa cliente: Distrito Urbano — comercio minorista de moda urbana con ventas en línea")
    set_run(run, italic=True, size=10.5, color=GRAY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    run = p.add_run("Documento de referencia para entendimiento del producto, diseño, construcción, pruebas y validación.")
    set_run(run, italic=True, size=10, color=GRAY)

    doc.add_page_break()


# =========================================================================
# DATA – Requisitos
# =========================================================================

RF_DATA = [
    ("RF-01", "Gestión de cuentas y control de acceso",
     "El sistema debe permitir la gestión de cuentas, autenticación y control de acceso de clientes y administradores.",
     "Cliente / Administrador", "Must",
     "Yo como Product Owner, para controlar el acceso seguro de clientes y administradores a la tienda digital, "
     "necesito que el sistema permita crear cuentas, iniciar sesión, cerrar sesión, actualizar perfiles administrativos "
     "y aplicar control de acceso basado en roles.",
     [
         "Crear formulario de registro de cliente con nombre, correo, teléfono y contraseña.",
         "Implementar inicio y cierre de sesión con validación de credenciales.",
         "Implementar control de acceso por rol (cliente vs administrador).",
         "Crear módulo de perfil administrativo.",
         "Validar seguridad básica: correo único, contraseña fuerte, acceso restringido.",
     ]),
    ("RF-02", "Vitrina pública con identidad visual",
     "El sistema debe mostrar la tienda digital, categorías y productos activos con identidad visual propia de la empresa.",
     "Cliente", "Must",
     "Yo como Product Owner, para ofrecer una experiencia visual clara y coherente con la marca, "
     "necesito que el sistema muestre la tienda digital con logo, colores, mensajes informativos, categorías y productos activos.",
     [
         "Diseñar interfaz principal con estructura visual y navegación.",
         "Integrar logo, colores corporativos, banner y mensajes visibles.",
         "Listar categorías activas del catálogo.",
         "Listar productos activos con nombre, imágenes, precio y disponibilidad.",
         "Aplicar estado de publicación: solo se visualizan productos y categorías activas.",
     ]),
    ("RF-03", "Búsqueda, filtros y ficha de producto",
     "El sistema debe permitir buscar, filtrar y consultar el detalle de productos con sus variantes.",
     "Cliente", "Must",
     "Yo como Product Owner, para facilitar que el cliente encuentre y revise productos antes de comprar, "
     "necesito que el sistema permita buscar, filtrar y consultar la ficha completa de productos con descripción, "
     "galería, variantes, precio, stock visible y valoración promedio.",
     [
         "Implementar buscador de productos por texto.",
         "Implementar filtros por categoría, rango de precio, color, talla y disponibilidad.",
         "Crear ficha de producto con descripción, galería, precio, stock y rating.",
         "Mostrar variantes (color y talla) con disponibilidad por combinación.",
         "Validar selección de variante específica antes de agregar al carrito.",
     ]),
    ("RF-04", "Carrito y captura de datos de checkout",
     "El sistema debe permitir gestionar el carrito de compras, validar stock y capturar datos para finalizar el pedido.",
     "Cliente", "Must",
     "Yo como Product Owner, para permitir que el cliente prepare correctamente su compra antes del pago, "
     "necesito que el sistema permita agregar productos al carrito, modificar cantidades, eliminar ítems, "
     "validar stock y capturar datos de entrega, facturación y contacto.",
     [
         "Implementar carrito de compras con la variante seleccionada.",
         "Modificar cantidades del carrito antes de confirmar la compra.",
         "Eliminar productos del carrito.",
         "Validar disponibilidad de stock antes de iniciar el checkout.",
         "Capturar datos de entrega, facturación y contacto.",
     ]),
    ("RF-05", "Cálculo, pasarela de pago y creación del pedido",
     "El sistema debe calcular valores de compra, integrar pasarela de pago y crear pedidos con estado actualizado.",
     "Sistema / Pasarela", "Must",
     "Yo como Product Owner, para procesar correctamente la compra, necesito que el sistema calcule subtotal, "
     "costos adicionales, descuentos, total a pagar, integre una pasarela de pago externa y cree el pedido con "
     "un identificador único y estado según el resultado del pago.",
     [
         "Calcular subtotal, costos adicionales, descuentos y total a pagar.",
         "Mostrar resumen de pago al cliente antes de confirmar.",
         "Integrar pasarela de pago externa (autorizar, rechazar o dejar pendiente).",
         "Crear pedido con identificador único cuando el pago sea aprobado.",
         "Validar respuesta de la pasarela: impedir confirmación sin respuesta válida.",
     ]),
    ("RF-06", "Consulta, actualización y notificación de pedidos",
     "El sistema debe permitir consultar, actualizar y notificar el estado de los pedidos.",
     "Cliente / Administrador / Sistema", "Must",
     "Yo como Product Owner, para permitir el seguimiento completo de las compras, necesito que el sistema permita "
     "al cliente consultar su historial y detalle de pedidos, manejar estados del pedido, permitir actualizaciones "
     "operativas por parte del administrador y notificar cambios relevantes.",
     [
         "Crear historial de pedidos del cliente.",
         "Mostrar detalle del pedido con productos, valores, datos de entrega y estado actual.",
         "Implementar estados: PAID, EN_PREPARACION, ENVIADO, ENTREGADO, CANCELADA.",
         "Permitir actualización administrativa del estado operativo.",
         "Notificar al cliente cambios relevantes (in-app + correo SMTP).",
     ]),
    ("RF-07", "Gestión de catálogo, variantes e inventario",
     "El sistema debe permitir la gestión administrativa de categorías, productos, variantes e inventario.",
     "Administrador", "Must",
     "Yo como Product Owner, para administrar el catálogo y el inventario, necesito que el sistema permita "
     "crear, consultar, editar, archivar y reactivar categorías y productos, definir variantes, asignar SKU, "
     "costo, precio y stock, registrar movimientos de inventario y generar alertas de stock mínimo.",
     [
         "Gestión de categorías (CRUD + archivado).",
         "Gestión de productos (CRUD + archivado + imágenes + descripciones).",
         "Definir variantes por color y talla con SKU único por combinación.",
         "Asignar costo, precio y stock independiente a cada variante.",
         "Registrar entradas, salidas y ajustes con motivo y responsable; alertas de stock mínimo.",
     ]),
    ("RF-08", "Administración financiera y reportes",
     "El sistema debe permitir gestionar información administrativa, financiera y reportes del negocio.",
     "Administrador", "Must",
     "Yo como Product Owner, para controlar la operación financiera del negocio, necesito que el sistema permita "
     "registrar empleados, costos operativos, gastos generales, calcular indicadores financieros, mostrar dashboards "
     "y exportar reportes administrativos y financieros.",
     [
         "Registrar empleados con datos básicos, cargo, salario y estado laboral.",
         "Registrar costos operativos y gastos generales.",
         "Calcular ventas brutas, COGS, costos operativos, nómina y utilidad neta por período.",
         "Dashboards con ventas por período, productos más vendidos y resumen financiero.",
         "Exportar reportes en PDF o CSV.",
     ]),
    ("RF-09", "Configuración de tienda, mensajes y reseñas",
     "El sistema debe permitir la configuración general de la tienda, mensajes informativos, parámetros globales y reseñas de productos.",
     "Administrador / Cliente", "Must",
     "Yo como Product Owner, para administrar la configuración visible y operativa de la tienda, "
     "necesito que el sistema permita configurar datos generales, administrar mensajes informativos, "
     "definir parámetros globales y permitir reseñas únicamente cuando exista una compra entregada.",
     [
         "Configurar nombre comercial, logo, colores, banner y datos de contacto.",
         "Administrar mensajes informativos (horarios, políticas, avisos).",
         "Definir parámetros globales (moneda, umbral de stock, estados habilitados).",
         "Implementar reseñas y valoraciones para productos comprados.",
         "Validar compra entregada antes de permitir la reseña + moderación admin.",
     ]),
]


RNF_DATA = [
    ("RNF-01", "Rendimiento, disponibilidad y concurrencia",
     "El sistema debe garantizar rendimiento, disponibilidad y capacidad concurrente bajo carga normal.",
     "Must",
     "Yo como Product Owner, para asegurar que la tienda digital funcione estable y rápida durante su operación normal, "
     "necesito que el sistema responda en tiempos adecuados, soporte usuarios concurrentes y mantenga una disponibilidad mensual aceptable.",
     [
         "95% de las consultas de catálogo y panel responden ≤ 3 segundos bajo carga normal.",
         "Soportar ≥ 300 clientes concurrentes.",
         "Soportar ≥ 20 sesiones administrativas concurrentes.",
         "Disponibilidad mensual ≥ 99,5% excluyendo mantenimientos programados.",
         "Optimizar consultas, endpoints y carga de información para cumplir tiempos definidos.",
     ]),
    ("RNF-02", "Seguridad, privacidad e integridad",
     "El sistema debe garantizar seguridad, privacidad e integridad de la información y las transacciones.",
     "Must",
     "Yo como Product Owner, para proteger datos personales, credenciales, operaciones administrativas y transacciones, "
     "necesito que el sistema aplique comunicación segura, almacenamiento protegido de contraseñas, control de acceso, "
     "mínimo privilegio e integridad transaccional.",
     [
         "Comunicación HTTPS/TLS entre todos los componentes y servicios externos.",
         "Contraseñas con hash bcrypt; expiración de sesiones por inactividad.",
         "Funciones administrativas restringidas a usuarios con rol admin (JWT validado).",
         "Principio de mínimo privilegio: cada microservicio accede solo a su esquema MySQL.",
         "Integridad transaccional: SAGA orquestada con compensaciones; sin doble descuento "
         "de inventario; pedidos solo PAID con confirmación válida de pasarela.",
     ]),
    ("RNF-03", "Usabilidad, mantenibilidad, respaldo y trazabilidad",
     "El sistema debe garantizar usabilidad, compatibilidad, mantenibilidad, respaldo, trazabilidad y observabilidad.",
     "Must",
     "Yo como Product Owner, para asegurar que la solución sea usable, mantenible, recuperable y diagnosticable, "
     "necesito que el sistema sea compatible con navegadores modernos, tenga interfaz accesible, "
     "arquitectura modular, respaldo, trazabilidad y registros suficientes para diagnosticar fallos.",
     [
         "Usuario nuevo puede completar primera compra en ≤ 5 minutos.",
         "Frontend opera en navegadores modernos y se adapta a móvil/escritorio.",
         "Backend Python y frontend React desacoplados, comunicados por REST JSON; MySQL como motor principal.",
         "Respaldos diarios con RPO ≤ 24 horas y RTO ≤ 4 horas.",
         "Trazabilidad: operaciones críticas registradas con actor, fecha, valor previo/nuevo + "
         "correlation-id propagado a través de los microservicios.",
     ]),
]


# =========================================================================
# Casos de uso ÚNICOS de la idea de la aplicación
# =========================================================================

CASOS_DE_USO = [
    {
        "id": "CU-01",
        "nombre": "Realizar checkout coordinando inventario y pasarela (SAGA)",
        "actor": "Cliente (principal); Pasarela de pago (secundario)",
        "rf": "RF-04, RF-05, RF-06",
        "objetivo":
            "Permitir que el cliente convierta el carrito en un pedido pagado, garantizando que NO "
            "se descuente inventario sin confirmación válida de la pasarela y que NO se cobre al "
            "cliente sin garantizar el stock. Es el caso de uso más crítico del sistema porque "
            "coordina tres bounded contexts (inventory, payment, commerce) bajo el patrón SAGA "
            "orquestada síncrona con compensaciones HTTP.",
        "precondicion":
            "El cliente está autenticado, tiene al menos un ítem en su carrito con stock disponible "
            "y diligenció los datos de entrega. El servicio Commerce orquesta la SAGA; el Inventory "
            "Service y el Payment Service están saludables (o el Circuit Breaker permite intentar).",
        "postcondicion_exito":
            "Se persiste una Order en estado PAID en commerce_db con OrderItems que incluyen "
            "unit_cost (snapshot del costo desde Inventory). El stock físico de cada variante se "
            "descuenta. Se notifica al cliente vía push in-app y correo SMTP. El carrito queda en "
            "estado checked_out.",
        "postcondicion_fallo":
            "NO se crea Order. La reserva de stock se libera automáticamente (compensación). El "
            "intento queda en FailedCheckoutAttempt con reason_code (out_of_stock | payment_rejected | "
            "payment_unavailable). El cliente recibe el error con detalle y su carrito queda intacto.",
        "flujo": [
            "El cliente hace clic en 'Confirmar pago' desde la pantalla de checkout.",
            "El Frontend envía POST /api/checkout con Idempotency-Key + datos de entrega.",
            "El API Gateway propaga la solicitud al Commerce Service con correlation_id.",
            "Commerce calcula subtotal y total tomando snapshots del carrito.",
            "Commerce llama POST /reserve a Inventory pidiendo reserva por 15 minutos para cada item.",
            "Inventory adquiere lock distribuido Redis SET NX EX por variante y hace SELECT FOR UPDATE en MySQL.",
            "Inventory crea StockReservation(status=PENDING) y devuelve 201 con reservation_ids.",
            "Commerce llama POST /payments a Payment Service con order_code, monto y card_token.",
            "Payment verifica que el Circuit Breaker esté CLOSED y llama POST /charge al payment-mock.",
            "El payment-mock devuelve {status: APPROVED, transaction_reference: AUTH-xxx}.",
            "Payment responde 200 OK a Commerce con el resultado APPROVED.",
            "Commerce llama POST /confirm/{order_id} a Inventory para descontar stock real.",
            "Inventory confirma: stock -= qty, reserved_stock -= qty y marca la reserva CONFIRMED.",
            "Commerce solicita /variants/by-ids a Inventory para obtener unit_cost de cada variante.",
            "Commerce persiste Order(status=PAID) + OrderItems(con unit_cost snapshot) + StatusHistory + AuditLog.",
            "Commerce inserta una Notification para el cliente y envía correo vía SMTP Mailhog.",
            "Commerce marca el carrito como checked_out.",
            "Commerce responde 201 Created con {order_id, order_code, status: PAID, total} al Gateway.",
            "El Frontend muestra la pantalla '¡Pago aprobado!' con el código del pedido y opciones de seguimiento.",
        ],
        "alternativos": [
            ("FA-01", "Stock insuficiente (sin variantes con stock)",
             [
                 "En el paso 6, Inventory detecta que alguna variante NO tiene stock suficiente "
                 "(considerando reservas activas).",
                 "Inventory devuelve 409 Conflict con la lista detallada de variantes sin stock.",
                 "Commerce registra el intento en FailedCheckoutAttempt(reason_code='out_of_stock').",
                 "Commerce NO crea Order y devuelve HTTP 409 al cliente con el detalle.",
                 "El Frontend muestra la pantalla de error con la lista de variantes agotadas y "
                 "ofrece 'Intentar de nuevo' o 'Volver al carrito'."
             ]),
            ("FA-02", "Pago rechazado por la pasarela",
             [
                 "En el paso 10, la pasarela responde {status: REJECTED, message: 'Fondos insuficientes'}.",
                 "Payment devuelve 200 OK con status=REJECTED a Commerce.",
                 "Commerce dispara la COMPENSACIÓN: POST /release a Inventory para liberar la reserva.",
                 "Inventory marca las reservas como RELEASED y libera el reserved_stock.",
                 "Commerce registra el intento en FailedCheckoutAttempt(reason_code='payment_rejected').",
                 "Commerce envía notificación + correo informando al cliente del rechazo.",
                 "Commerce devuelve HTTP 402 con {code: 'payment_rejected', message: 'Fondos insuficientes'}.",
                 "El Frontend muestra la pantalla 'Pago rechazado' con razón y botón 'Intentar con otro método'."
             ]),
            ("FA-03", "Circuit Breaker abierto (pasarela caída)",
             [
                 "Antes del paso 9, Payment Service detecta que el Circuit Breaker está en estado OPEN "
                 "(5 fallos consecutivos en últimos 60 segundos).",
                 "Payment responde inmediatamente 503 Service Unavailable a Commerce SIN tocar la pasarela.",
                 "Commerce ejecuta la compensación: POST /release a Inventory.",
                 "Commerce registra FailedCheckoutAttempt(reason_code='payment_unavailable').",
                 "Commerce devuelve 503 al cliente con {code: 'payment_unavailable'}.",
                 "El Frontend muestra 'La pasarela está temporalmente fuera de servicio. Tu carrito "
                 "está intacto, intenta en unos minutos.'"
             ]),
            ("FA-04", "Reintento idempotente",
             [
                 "Si el cliente hace doble clic o reintenta el envío con la misma Idempotency-Key, "
                 "Commerce detecta que ya existe una Order con ese correlation_id.",
                 "Commerce devuelve directamente la Order ya creada sin reprocesar.",
                 "NO hay doble cobro ni doble descuento de stock."
             ]),
        ],
    },
    {
        "id": "CU-02",
        "nombre": "Reseñar producto entregado con moderación administrativa",
        "actor": "Cliente (principal); Administrador (moderador)",
        "rf": "RF-09",
        "objetivo":
            "Permitir que el cliente publique una opinión calificada (1–5 estrellas + comentario) "
            "ÚNICAMENTE sobre productos que efectivamente compró y recibió. La reseña entra como "
            "pendiente, requiere aprobación del administrador y, una vez publicada, recalcula el "
            "rating promedio del producto en el Catalog Service. Es único del dominio porque combina "
            "validación de pedido entregado + moderación + propagación cross-context del rating.",
        "precondicion":
            "El cliente está autenticado. Existe al menos un pedido del cliente en estado ENTREGADO "
            "que contiene el producto a reseñar. El cliente NO ha enviado previamente una reseña "
            "para esa combinación (producto, pedido).",
        "postcondicion_exito":
            "Se crea Review(approved=False) en commerce_db. La reseña queda pendiente de aprobación. "
            "El cliente recibe confirmación. En el detalle del pedido el producto se marca como "
            "'Reseñado' (badge verde). Cuando el administrador aprueba: el flag approved cambia a True, "
            "Commerce recalcula el promedio y notifica al Catalog Service que actualiza el "
            "RatingSummary del producto.",
        "postcondicion_fallo":
            "No se crea reseña. El cliente ve un mensaje claro: 'Solo puedes reseñar productos "
            "comprados y entregados' o 'Ya reseñaste este producto para este pedido'.",
        "flujo": [
            "El cliente abre /mis-pedidos y selecciona un pedido en estado ENTREGADO.",
            "En el detalle del pedido, por cada producto comprado aparece un botón 'Dejar reseña'.",
            "El cliente hace clic en 'Dejar reseña' del producto X.",
            "El Frontend navega a /resenas/{orderId}/{productId}.",
            "El cliente selecciona una calificación (1–5 estrellas) con el star picker.",
            "El cliente escribe un comentario en el textarea (con contador de caracteres).",
            "El cliente hace clic en 'Publicar reseña'.",
            "El Frontend envía POST /api/reviews {order_id, product_id, rating, comment} al Gateway.",
            "Commerce Service valida que (order_id, user_id) corresponda a un pedido ENTREGADO del cliente.",
            "Commerce verifica que el product_id esté en los OrderItems del pedido.",
            "Commerce verifica que NO exista ya una Review para (user_id, product_id, order_id).",
            "Commerce persiste Review(approved=False) en commerce_db.",
            "Commerce responde 201 Created con la reseña creada.",
            "El Frontend muestra toast: 'Reseña publicada. Quedará pendiente de aprobación.'",
            "El Frontend marca el producto en el pedido con badge 'Reseñado' (verde).",
            "// — momento posterior, asíncrono —",
            "El administrador entra a /admin/resenas y ve la nueva reseña en el filtro 'Pendientes'.",
            "El administrador revisa el contenido (rating, comentario, producto, pedido, cliente).",
            "El administrador hace clic en el botón 'Aprobar'.",
            "El Frontend envía PATCH /api/admin/reviews/{id}/approve.",
            "Commerce marca approved=True, recalcula el promedio de todas las reseñas aprobadas del producto.",
            "Commerce llama PUT /admin/products/{id}/rating al Catalog Service con {average, count}.",
            "Catalog actualiza el RatingSummary del producto e invalida el caché Redis del producto.",
            "El producto aparece con el nuevo rating en la ficha pública con el sello 'Compra verificada'.",
        ],
        "alternativos": [
            ("FA-01", "Producto no comprado o no entregado",
             [
                 "Si el cliente intenta enviar una reseña de un producto que no compró o cuyo pedido "
                 "no está en estado ENTREGADO, Commerce devuelve 409 Conflict.",
                 "Mensaje: 'Solo puedes reseñar productos comprados y entregados.'",
                 "El Frontend muestra el error en un toast y no permite reenviar."
             ]),
            ("FA-02", "Reseña duplicada",
             [
                 "Si el cliente ya envió una reseña para esa combinación (producto, pedido), Commerce "
                 "devuelve 409 Conflict con 'Ya reseñaste este producto para este pedido'."
             ]),
            ("FA-03", "Administrador rechaza la reseña",
             [
                 "Si el contenido no cumple las políticas (insultos, spam, contenido off-topic), el "
                 "administrador hace clic en 'Rechazar/Eliminar'.",
                 "El Frontend envía DELETE /api/admin/reviews/{id}.",
                 "Commerce elimina la reseña. Si estaba aprobada, recalcula el rating descontándola y "
                 "notifica a Catalog.",
                 "Si estaba pendiente, simplemente se elimina sin afectar el rating."
             ]),
        ],
    },
    {
        "id": "CU-03",
        "nombre": "Configurar variantes Nike-style (color + talla) con stock por combinación",
        "actor": "Administrador",
        "rf": "RF-07",
        "objetivo":
            "Permitir al administrador definir variantes de un producto cruzando dos dimensiones "
            "(color y talla) de forma que cada COMBINACIÓN única tenga su propio SKU, costo, precio "
            "y stock. Es único del dominio porque implementa el modelo de variantes estilo Nike "
            "(elegir color → ver tallas disponibles), con visualización real del color (hex) y "
            "restricción de unicidad en la combinación.",
        "precondicion":
            "El administrador está autenticado con rol admin. El producto base ya existe en el "
            "Catalog Service (con precio base, descripción e imágenes). El Inventory Service está "
            "disponible.",
        "postcondicion_exito":
            "Se crea una nueva ProductVariant en inventory_db con (product_id, color, color_hex, "
            "size, sku, cost, price, stock). La combinación (product_id, color, size) queda registrada "
            "como única (UniqueConstraint uq_variant_combo). El cliente, al entrar al detalle del "
            "producto, ve el selector de color con la bolita del hex real; al elegir color ve las "
            "tallas disponibles para ese color con su stock.",
        "postcondicion_fallo":
            "No se crea la variante. El sistema muestra mensaje explicando la razón: 'Ya existe una "
            "variante con color X y talla Y para este producto' o 'SKU duplicado'.",
        "flujo": [
            "El administrador entra al panel admin → Catálogo → Variantes.",
            "Selecciona el producto al que quiere agregar la variante.",
            "El sistema muestra el formulario con: SKU (autogenerable), Color (nombre + picker hex "
            "+ 12 presets), Talla, Costo unitario, Precio de venta, Stock inicial.",
            "El sistema muestra el precio base del producto como referencia.",
            "El administrador escribe el nombre del color (ej. 'Negro').",
            "Hace clic en uno de los 12 presets de color (o usa el input type=color) para seleccionar "
            "el hex visual (#111111).",
            "Selecciona o escribe la talla (ej. 'M').",
            "Ingresa el costo unitario (ej. 22000).",
            "Ingresa el precio de venta (ej. 49000).",
            "El sistema calcula y muestra en vivo el margen estimado: $27.000 (55,1%) en verde.",
            "Ingresa el stock inicial (ej. 15 unidades).",
            "Hace clic en 'Crear variante'.",
            "El Frontend valida en cliente que NO exista ya una variante activa con (color, talla) idéntica.",
            "El Frontend envía POST /api/admin/inventory/variants al Gateway.",
            "El Gateway propaga al Inventory Service con el JWT admin.",
            "Inventory valida que el product_id exista en Catalog (llamada cross-service).",
            "Inventory valida que el SKU no esté duplicado.",
            "Inventory valida que la combinación (product_id, color, size) sea única.",
            "Inventory persiste ProductVariant con color_hex incluido.",
            "Inventory devuelve 201 Created con la variante creada.",
            "El Frontend agrega la variante a la tabla con la bolita del color visible.",
            "El cliente, al entrar al detalle del producto, ve el selector Nike-style: las bolitas "
            "de color (agrupadas), y al elegir un color ve las tallas disponibles con su stock.",
        ],
        "alternativos": [
            ("FA-01", "Combinación duplicada (mismo color + misma talla)",
             [
                 "En el paso 13 (validación cliente) o paso 18 (validación servidor), si ya existe una "
                 "variante con la misma combinación, Inventory devuelve 409 Conflict.",
                 "Mensaje: 'Ya existe una variante para producto X con color=Negro y talla=M (SKU CAM-NEG-M). "
                 "Edítala en vez de crear una nueva.'",
                 "El Frontend muestra el toast y NO crea la variante."
             ]),
            ("FA-02", "Producto no existe en Catalog",
             [
                 "Si el product_id no existe en Catalog (caso raro, normalmente solo si el producto "
                 "fue eliminado físicamente), Inventory devuelve 422 Unprocessable Entity.",
                 "Mensaje: 'product_id=X no existe en Catalog.'"
             ]),
            ("FA-03", "Inventory caído al validar",
             [
                 "Si Catalog Service está temporalmente caído cuando Inventory intenta validar la "
                 "existencia del producto, Inventory entra en modo degradado y permite la creación "
                 "(con un log de WARN para auditoría posterior).",
                 "Esto evita bloquear operaciones admin críticas durante una caída transitoria."
             ]),
            ("FA-04", "SKU duplicado",
             [
                 "Si el administrador escribió un SKU que ya existe globalmente, Inventory devuelve "
                 "409 con 'Ya existe una variante con SKU=X'. El admin puede dejar el SKU vacío para "
                 "que se autogenere (P{product_id}-{timestamp}) y evitar colisiones."
             ]),
        ],
    },
    {
        "id": "CU-04",
        "nombre": "Operar transición logística del pedido con notificación al cliente",
        "actor": "Administrador",
        "rf": "RF-06",
        "objetivo":
            "Permitir al administrador hacer avanzar un pedido a través de su ciclo logístico "
            "(PAID → EN_PREPARACION → ENVIADO → ENTREGADO) o cancelarlo, garantizando que cada "
            "transición sea válida según la máquina de estados, quede en historia con actor y "
            "timestamp, dispare una notificación in-app + correo al cliente y se persista en "
            "OrderAuditLog para trazabilidad. Es único porque combina validación de máquina de "
            "estados + auditoría con correlation_id + notificación multi-canal.",
        "precondicion":
            "El administrador está autenticado con rol admin. Existe al menos un pedido en estado "
            "PAID, EN_PREPARACION o ENVIADO en el sistema.",
        "postcondicion_exito":
            "Order.status actualizado a la nueva etapa. Se inserta OrderStatusHistory(from_status, "
            "to_status, changed_by=admin_id, changed_at, notes). Se inserta OrderAuditLog con la "
            "acción. Se crea Notification para el cliente. Se envía correo SMTP al cliente con el "
            "cambio. El cliente ve el progreso actualizado en el timeline de su pedido.",
        "postcondicion_fallo":
            "No se actualiza el estado. El administrador recibe error claro con las transiciones "
            "permitidas desde el estado actual.",
        "flujo": [
            "El administrador entra al panel admin → Pedidos.",
            "El sistema muestra solo pedidos en estados operativos (PAID, EN_PREPARACION, ENVIADO, "
            "ENTREGADO, CANCELADA). Los rechazados/sin stock NO aparecen aquí (están en la bitácora).",
            "El administrador filtra por estado 'Pagado' y selecciona un pedido.",
            "Hace clic en el dropdown de estado y selecciona 'En preparación'.",
            "El Frontend envía PATCH /api/admin/orders/{id}/status {new_status: 'EN_PREPARACION'} al Gateway.",
            "Commerce valida que la transición PAID → EN_PREPARACION sea legal según la máquina de estados.",
            "Commerce actualiza Order.status = 'EN_PREPARACION'.",
            "Commerce inserta una entrada en OrderStatusHistory con from_status=PAID, to_status=EN_PREPARACION, "
            "changed_by=admin_id, notes='Cambio administrativo', changed_at=now().",
            "Commerce inserta evento en OrderAuditLog con action='status_change_PAID_to_EN_PREPARACION' "
            "y correlation_id propagado.",
            "Commerce crea Notification(user_id=cliente, title='Tu pedido está en preparación', "
            "message='El pedido ORD-XXX ahora está EN_PREPARACION').",
            "Commerce envía un correo SMTP a Mailhog con el mismo contenido al email del cliente.",
            "Commerce devuelve 200 OK con la nueva versión del pedido.",
            "El Frontend refresca la tabla y muestra un toast de confirmación.",
            "El cliente, al refrescar /mis-pedidos, ve el badge actualizado y un nuevo nodo activo "
            "en el timeline 'En preparación' con la fecha.",
            "El administrador puede repetir el flujo para EN_PREPARACION → ENVIADO → ENTREGADO.",
        ],
        "alternativos": [
            ("FA-01", "Transición inválida",
             [
                 "Si el administrador intenta retroceder (ej. ENVIADO → PAID) o saltar etapas "
                 "(ej. PAID → ENVIADO), Commerce devuelve 409 Conflict.",
                 "Mensaje: 'Transición inválida: PAID → ENVIADO. Permitidas desde PAID: "
                 "EN_PREPARACION, CANCELADA.'",
                 "El Frontend muestra un toast de error y mantiene el estado anterior."
             ]),
            ("FA-02", "Cancelación administrativa",
             [
                 "Desde PAID o EN_PREPARACION el administrador puede transicionar a CANCELADA.",
                 "El sistema dispara el mismo flujo: history + audit + notification + correo con "
                 "título 'Tu pedido fue cancelado'.",
                 "Los estados ENVIADO y ENTREGADO no admiten cancelación (el producto ya salió)."
             ]),
            ("FA-03", "Caída del servidor SMTP",
             [
                 "Si Mailhog (o el servidor SMTP real) no responde, el correo no se envía pero el "
                 "cambio de estado SÍ se persiste y la notificación in-app SÍ se crea.",
                 "El fallo del envío de correo se registra en logs para reintento manual."
             ]),
        ],
    },
    {
        "id": "CU-05",
        "nombre": "Consultar dashboard financiero con COGS y márgenes reales",
        "actor": "Administrador",
        "rf": "RF-08",
        "objetivo":
            "Permitir al administrador consultar los indicadores financieros del negocio (ventas, "
            "COGS, margen bruto, gastos, nómina, utilidad neta) en un período seleccionado, con "
            "evolución temporal por día/mes/año y exportación a PDF/CSV. Es único del dominio porque "
            "calcula COGS REAL gracias al snapshot de unit_cost capturado en cada OrderItem al "
            "momento del checkout, no a costos actuales (que podrían haber cambiado).",
        "precondicion":
            "El administrador está autenticado con rol admin. Existe al menos un pedido en estado "
            "PAID o posterior. El Backend ha capturado unit_cost en OrderItem al hacer checkout "
            "(o se ejecutó el backfill desde Inventory para pedidos históricos).",
        "postcondicion_exito":
            "El administrador visualiza KPIs (ventas, COGS, margen bruto %, margen neto, ticket "
            "promedio), gráfica de líneas con la evolución del período, gráfica de torta con la "
            "distribución (COGS / gastos / nómina / utilidad), gráfica de barras con productos más "
            "vendidos y tabla con detalle por período. Si exporta: se descarga un CSV con BOM UTF-8 "
            "o se abre una ventana HTML imprimible 'reporte premium' con A4, hero, KPIs y tablas.",
        "postcondicion_fallo":
            "Si Commerce no responde o no hay pedidos, el dashboard muestra ceros y un mensaje "
            "'Sin datos para este período'.",
        "flujo": [
            "El administrador entra al menú 'Finanzas' del panel.",
            "El Frontend solicita GET /api/admin/finance/summary?granularity=month al Gateway.",
            "Commerce calcula los indicadores consultando commerce_db:",
            "  • ventas_brutas = SUM(Order.total) WHERE status IN (PAID, EN_PREPARACION, ENVIADO, ENTREGADO)",
            "  • cogs = SUM(OrderItem.unit_cost × quantity) — usa el snapshot de costo capturado al checkout",
            "  • margen_bruto = ventas_brutas - cogs",
            "  • gastos_operativos = SUM(Expense.amount) en el período",
            "  • nomina = SUM(Employee.salary) WHERE employment_status='active'",
            "  • utilidad_neta = margen_bruto - gastos_operativos - nomina",
            "  • timeseries: misma agregación pero agrupada por día / mes / año.",
            "Commerce responde con el JSON estructurado del summary.",
            "El Frontend renderiza 4 KPIs (ventas, COGS, margen bruto, utilidad neta).",
            "Renderiza LineChart con tres series: ventas, COGS, margen bruto sobre el tiempo.",
            "Renderiza PieChart con la distribución financiera del período.",
            "Renderiza BarChart con productos más vendidos por unidades.",
            "Renderiza dos tablas: detalle por período y gastos por tipo.",
            "El administrador cambia el filtro a 'Año' o ajusta las fechas manualmente.",
            "El Frontend re-invoca el endpoint con los nuevos parámetros y refresca todas las gráficas.",
            "El administrador hace clic en 'Reporte PDF'.",
            "El Frontend abre una ventana HTML imprimible con: hero corporativo + KPIs + tabla + footer.",
            "El administrador imprime o guarda como PDF desde el diálogo del navegador.",
        ],
        "alternativos": [
            ("FA-01", "Exportar CSV",
             [
                 "Si el administrador hace clic en 'CSV' en vez de 'PDF', el Frontend genera un Blob "
                 "con BOM UTF-8 que incluye KPIs, timeseries, productos más vendidos y gastos por tipo, "
                 "y dispara la descarga del archivo .csv compatible con Excel."
             ]),
            ("FA-02", "Utilidad neta negativa",
             [
                 "Si la utilidad neta calculada es negativa (los costos + gastos + nómina superan las "
                 "ventas brutas), el Frontend muestra una alerta roja en la parte superior: 'Utilidad "
                 "neta negativa: los costos superan las ventas. Revisa gastos y nómina.'",
                 "El reporte PDF muestra un badge rojo 'En pérdidas' en la sección de resultado operativo."
             ]),
            ("FA-03", "Pedidos históricos sin unit_cost",
             [
                 "Si la base contiene pedidos creados antes de la captura de unit_cost, sus OrderItems "
                 "tienen unit_cost=0 y el COGS de esos períodos sale subestimado.",
                 "El administrador ejecuta POST /api/admin/maintenance/backfill-costs para que Commerce "
                 "recorra los OrderItems con unit_cost=0 y los rellene consultando a Inventory por batch.",
                 "Tras el backfill, los KPIs reflejan COGS correcto."
             ]),
        ],
    },
]


# =========================================================================
# Diagramas PlantUML (todos los códigos completos)
# =========================================================================

PLANTUML_CASOS_USO = r"""@startuml DiagramaCasosUso
left to right direction
skinparam packageStyle rectangle
skinparam usecase {
  BackgroundColor #DCE6F1
  BorderColor #1F497D
  ArrowColor #1F497D
  FontColor #172026
}
skinparam actor {
  BackgroundColor #FFFFFF
  BorderColor #1F497D
  FontColor #172026
}
skinparam shadowing false
title <b>Diagrama de Casos de Uso — Distrito Urbano</b>\n<i>Tienda Digital con arquitectura de microservicios</i>

actor "Cliente" as cliente
actor "Administrador" as admin
actor "Pasarela de Pago" as pasarela <<sistema externo>>
actor "Servicio SMTP" as smtp <<sistema externo>>

rectangle "Sistema: Tienda Digital — Distrito Urbano" {

  ' ───── Casos de uso del Cliente ─────
  usecase "Registrarse" as UC1
  usecase "Iniciar sesión" as UC2
  usecase "Explorar catálogo" as UC3
  usecase "Buscar producto" as UC3a
  usecase "Consultar detalle\nde producto" as UC4
  usecase "Gestionar carrito" as UC5
  usecase "Realizar checkout\n(SAGA orquestada)" as UC6
  usecase "Procesar pago" as UC7
  usecase "Consultar pedidos" as UC8
  usecase "Reseñar producto\nentregado" as UC9
  usecase "Recibir notificación" as UC10

  ' ───── Casos de uso del Administrador ─────
  usecase "Gestionar catálogo\n(categorías, productos)" as UC11
  usecase "Configurar variantes\n(color + talla)" as UC12
  usecase "Operar pedidos\n(transición de estados)" as UC13
  usecase "Aprobar reseñas" as UC14
  usecase "Consultar dashboard\nfinanciero" as UC15
  usecase "Configurar tienda" as UC16
  usecase "Consultar bitácora\nde auditoría" as UC17
  usecase "Gestionar empleados\ny gastos" as UC18
}

' ───── Asociaciones del Cliente ─────
cliente --> UC1
cliente --> UC2
cliente --> UC3
cliente --> UC3a
cliente --> UC4
cliente --> UC5
cliente --> UC6
cliente --> UC8
cliente --> UC9
cliente --> UC10

' ───── Asociaciones del Administrador ─────
admin --> UC2
admin --> UC11
admin --> UC12
admin --> UC13
admin --> UC14
admin --> UC15
admin --> UC16
admin --> UC17
admin --> UC18

' ───── Asociaciones a sistemas externos ─────
UC7 --> pasarela : <<communicates>>
UC10 ..> smtp : <<communicates>>
UC13 ..> smtp : <<notifica al cliente>>

' ───── Relaciones <<include>> (obligatorias) ─────
UC6 ..> UC5 : <<include>>
UC6 ..> UC7 : <<include>>
UC9 ..> UC8 : <<include>>
UC14 ..> UC15 : <<include>>

' ───── Relaciones <<extend>> (opcionales) ─────
UC3a ..> UC3 : <<extend>>
UC10 ..> UC7 : <<extend>>

note right of UC6
  Caso de uso CENTRAL del sistema:
  coordina Inventory + Payment + Commerce
  con compensaciones HTTP si la SAGA falla.
end note

note bottom of UC9
  Único en el dominio: la reseña requiere
  un pedido ENTREGADO + moderación admin
  antes de publicarse y actualizar el rating.
end note

@enduml
"""


PLANTUML_CLASES = r"""@startuml DiagramaClases
skinparam classAttributeIconSize 0
skinparam classFontStyle bold
skinparam class {
  BackgroundColor #DCE6F1
  BorderColor #1F497D
  ArrowColor #1F497D
  FontColor #172026
}
skinparam shadowing false
hide circle
title <b>Diagrama de Clases — Modelo de dominio</b>\n<i>Agrupado por microservicio. Referencias cross-context son lógicas (no FK física).</i>

' ═══════════════ AUTH SERVICE ═══════════════
package "Auth Service (auth_db)" #E0F2FE {
  class User {
    + id: int <<PK>>
    + email: str <<unique>>
    + password_hash: str
    + name: str
    + phone: str
    + role: str
    + created_at: datetime
    --
    + verify_password(): bool
    + generate_jwt(): str
  }

  class RefreshToken {
    + id: int <<PK>>
    + user_id: int <<FK>>
    + token_hash: str
    + expires_at: datetime
    + revoked: bool
  }

  class AccessLog {
    + id: int <<PK>>
    + user_id: int <<FK, nullable>>
    + action: str
    + ip: str
    + user_agent: str
    + correlation_id: str
    + created_at: datetime
  }

  User "1" -- "0..*" RefreshToken
  User "1" -- "0..*" AccessLog
}

' ═══════════════ CATALOG SERVICE ═══════════════
package "Catalog Service (catalog_db)" #D4F0E1 {
  class Category {
    + id: int <<PK>>
    + name: str
    + description: str
    + active: bool
    + archived: bool
  }

  class Product {
    + id: int <<PK>>
    + category_id: int <<FK>>
    + name: str
    + description: str
    + long_description: text
    + base_price: decimal
    + image_url: str
    + published: bool
    + archived: bool
  }

  class ProductImage {
    + id: int <<PK>>
    + product_id: int <<FK>>
    + image_url: str
    + alt_text: str
  }

  class RatingSummary {
    + product_id: int <<PK>>
    + average: float
    + count: int
    + updated_at: datetime
  }

  class StoreSetting {
    + id: int <<PK>>
    + commercial_name: str
    + logo_url: str
    + primary_color: str
    + currency: str
  }

  Category "1" -- "0..*" Product
  Product "1" -- "0..*" ProductImage
  Product "1" -- "0..1" RatingSummary
}

' ═══════════════ INVENTORY SERVICE ═══════════════
package "Inventory Service (inventory_db)" #FED7AA {
  class ProductVariant {
    + id: int <<PK>>
    + product_id: int <<ref-logica>>
    + sku: str <<unique>>
    + color: str
    + color_hex: str
    + size: str
    + cost: decimal
    + price: decimal
    + stock: int
    + reserved_stock: int
    + active: bool
    --
    + available(): int
    --
    {static} uq_variant_combo(product_id, color, size)
  }

  class StockReservation {
    + id: int <<PK>>
    + variant_id: int <<FK>>
    + order_id: str <<ref-logica>>
    + quantity: int
    + status: str
    + expires_at: datetime
  }

  class StockMovement {
    + id: int <<PK>>
    + variant_id: int <<FK>>
    + movement_type: str
    + quantity: int
    + reason: str
    + user_id: int
    + correlation_id: str
  }

  class LowStockAlert {
    + id: int <<PK>>
    + variant_id: int <<FK>>
    + threshold: int
    + resolved: bool
  }

  ProductVariant "1" -- "0..*" StockReservation
  ProductVariant "1" -- "0..*" StockMovement
  ProductVariant "1" -- "0..*" LowStockAlert
}

' ═══════════════ COMMERCE SERVICE ═══════════════
package "Commerce Service (commerce_db)" #DBEAFE {
  class Cart {
    + id: int <<PK>>
    + user_id: int <<ref-logica>>
    + status: str
  }

  class CartItem {
    + id: int <<PK>>
    + cart_id: int <<FK>>
    + variant_id: int <<ref-logica>>
    + quantity: int
    + unit_price: decimal
  }

  class Order {
    + id: int <<PK>>
    + order_code: str <<unique>>
    + user_id: int <<ref-logica>>
    + status: str
    + payment_status: str
    + subtotal: decimal
    + total: decimal
    + delivery_name: str
    + delivery_address: str
    + correlation_id: str
  }

  class OrderItem {
    + id: int <<PK>>
    + order_id: int <<FK>>
    + variant_id: int <<ref-logica>>
    + product_id: int <<ref-logica>>
    + product_name: str
    + quantity: int
    + unit_price: decimal
    + unit_cost: decimal
    + total: decimal
  }

  class OrderStatusHistory {
    + id: int <<PK>>
    + order_id: int <<FK>>
    + from_status: str
    + to_status: str
    + changed_by: int
    + changed_at: datetime
  }

  class OrderAuditLog {
    + id: int <<PK>>
    + order_id: int <<FK, nullable>>
    + action: str
    + performed_by: int
    + correlation_id: str
  }

  class FailedCheckoutAttempt {
    + id: int <<PK>>
    + user_id: int
    + attempt_code: str
    + reason_code: str
    + correlation_id: str
  }

  class Review {
    + id: int <<PK>>
    + product_id: int <<ref-logica>>
    + order_id: int <<FK>>
    + user_id: int
    + rating: int
    + comment: str
    + approved: bool
  }

  class Notification {
    + id: int <<PK>>
    + user_id: int
    + order_id: int <<FK, nullable>>
    + title: str
    + read: bool
  }

  class Employee {
    + id: int <<PK>>
    + name: str
    + salary: decimal
    + employment_status: str
  }

  class Expense {
    + id: int <<PK>>
    + amount: decimal
    + type: str
    + expense_date: date
  }

  Cart "1" -- "0..*" CartItem
  Order "1" -- "1..*" OrderItem
  Order "1" -- "0..*" OrderStatusHistory
  Order "1" -- "0..*" OrderAuditLog
  Order "1" -- "0..*" Review
  Order "1" -- "0..*" Notification
}

' ═══════════════ PAYMENT SERVICE ═══════════════
package "Payment Service (payments_db)" #FEE2E2 {
  class Payment {
    + id: int <<PK>>
    + order_id: str <<ref-logica>>
    + amount: decimal
    + status: str
    + transaction_reference: str
  }
}

' ═══════════════ Referencias cross-context ═══════════════
User .[#777,dashed]. Order : "user_id (lógico)"
Product .[#777,dashed]. ProductVariant : "product_id (lógico)"
Product .[#777,dashed]. OrderItem : "product_id (lógico)"
ProductVariant .[#777,dashed]. CartItem : "variant_id (lógico)"
ProductVariant .[#777,dashed]. OrderItem : "variant_id (lógico)"
Order .[#777,dashed]. Payment : "order_id (lógico)"
Order .[#777,dashed]. StockReservation : "order_id (lógico)"

@enduml
"""


PLANTUML_SECUENCIA = r"""@startuml DiagramaSecuenciaCheckout
skinparam sequence {
  ParticipantBackgroundColor #DCE6F1
  ParticipantBorderColor #1F497D
  ParticipantFontStyle bold
  LifeLineBorderColor #1F497D
  ActorBackgroundColor #FFFFFF
  ArrowColor #1F497D
}
skinparam shadowing false
title <b>Diagrama de Secuencia — Realizar Checkout (SAGA orquestada síncrona)</b>\n<i>Flujo crítico que coordina los 5 microservicios</i>

actor "Cliente" as cliente
participant "Frontend\nReact" as front
participant "API Gateway\n(Nginx)" as gw
participant "Commerce\nService" as commerce
participant "Inventory\nService" as inventory
participant "Payment\nService" as payment
participant "Pasarela\nMock" as mock

== 1. Inicio del checkout ==
cliente -> front : Confirmar pago
front -> gw : POST /api/checkout\n[Idempotency-Key, datos entrega]
gw -> commerce : proxy_pass + correlation_id
activate commerce

commerce -> commerce : Calcular totales\n(snapshot del carrito)

== 2. Reservar inventario ==
commerce -> inventory : POST /reserve\n{order_code, items, ttl=900}
activate inventory
inventory -> inventory : Lock distribuido Redis\nSET NX EX por variante
inventory -> inventory : SELECT FOR UPDATE\nverifica stock disponible
inventory -> inventory : Crea StockReservation\n(status=PENDING)
inventory --> commerce : 201 {reservation_ids}
deactivate inventory

== 3. Cobrar a través de la pasarela ==
commerce -> payment : POST /payments\n{order_code, amount, card_token}
activate payment
payment -> payment : Verifica Circuit Breaker\n(CLOSED?)

alt CB == CLOSED
  payment -> mock : POST /charge
  activate mock
  mock --> payment : 200 {status: APPROVED,\ntransaction_ref}
  deactivate mock
  payment --> commerce : 200 {status: APPROVED}
else CB == OPEN
  payment --> commerce : 503 Service Unavailable
  commerce -> inventory : POST /release\n(compensación)
  commerce --> gw : 503 {code: payment_unavailable}
  gw --> front : 503
  front --> cliente : "Pasarela no disponible.\nTu carrito está intacto."
end
deactivate payment

== 4. Confirmar inventario (solo si APPROVED) ==
commerce -> inventory : POST /confirm/{order_id}
activate inventory
inventory -> inventory : stock -= qty\nreserved_stock -= qty
inventory -> inventory : Marca reserva\ncomo CONFIRMED
inventory --> commerce : 200 OK
deactivate inventory

== 5. Snapshot de costos para COGS ==
commerce -> inventory : GET /variants/by-ids
activate inventory
inventory --> commerce : [{variant_id, cost, price}]
deactivate inventory

== 6. Persistir orden y notificar ==
commerce -> commerce : Persiste Order(PAID)\n+ OrderItems(unit_cost)\n+ StatusHistory\n+ AuditLog
commerce -> commerce : Crea Notification + envía SMTP
commerce -> commerce : Cart.status = checked_out

commerce --> gw : 201 Created\n{order_id, order_code, status:PAID, total}
deactivate commerce
gw --> front : 201
front --> cliente : Mostrar "¡Pago aprobado!"

note over commerce, inventory
  <b>Política MVP:</b> la Order SOLO se persiste si la SAGA llega a PAID.
  Si REJECTED/SIN_STOCK/CB-open → compensación + FailedCheckoutAttempt + error HTTP.
  No hay "pedidos fantasma" en el panel admin.
end note

@enduml
"""


PLANTUML_ACTIVIDAD = r"""@startuml DiagramaActividad
skinparam shadowing false
skinparam activity {
  BackgroundColor #DCE6F1
  BorderColor #1F497D
  ArrowColor #1F497D
  DiamondBackgroundColor #FED7AA
  DiamondBorderColor #F59E0B
  FontColor #172026
}
title <b>Diagrama de Actividad — Proceso de compra del cliente</b>\n<i>Flujo end-to-end con decisiones</i>

start

:Iniciar sesión en la tienda;
:Explorar catálogo\no buscar producto;
:Abrir detalle del producto;
:Seleccionar variante\n(color → talla);

if (¿Variante tiene stock?) then (Sí)
  :Agregar al carrito;
  :Revisar carrito\ny ajustar cantidades;
  :Capturar datos\nde entrega y facturación;
  :Confirmar pago;

  fork
    :Frontend envía\nPOST /api/checkout;
  fork again
    :Backend ejecuta SAGA;
  end fork

  if (¿SAGA completa con éxito?) then (PAID)
    :Persistir Order(PAID)\n+ confirmar stock;
    :Notificar al cliente\n(in-app + correo);
    :Mostrar pantalla\n"¡Pago aprobado!";
  else (falla)
    if (¿Razón del fallo?) then (out_of_stock)
      :Mostrar lista de\nvariantes sin stock;
    elseif (payment_rejected) then
      :Mostrar "Pago rechazado"\ncon razón;
      :Liberar reserva\n(compensación);
    elseif (payment_unavailable) then
      :Mostrar "Pasarela\nno disponible";
      :Liberar reserva\n(compensación);
    endif
    :Registrar intento\nen FailedCheckoutAttempt;
    :Cliente puede\nreintentar;
  endif

else (No)
  :Mostrar AGOTADO;
  :Cliente vuelve\nal catálogo;
  detach
endif

stop

@enduml
"""


PLANTUML_ESTADO = r"""@startuml DiagramaEstado
skinparam shadowing false
skinparam state {
  BackgroundColor #DCE6F1
  BorderColor #1F497D
  FontColor #172026
  ArrowColor #1F497D
}
title <b>Diagrama de Estado — Ciclo de vida del Pedido</b>\n<i>Política MVP: la Order solo existe si el checkout llega a PAID</i>

[*] --> PAID : SAGA exitosa\n(reserve + charge + confirm)

PAID --> EN_PREPARACION : admin\ninicia preparación
PAID --> CANCELADA : admin cancela

EN_PREPARACION --> ENVIADO : admin\ndespacha
EN_PREPARACION --> CANCELADA : admin cancela

ENVIADO --> ENTREGADO : admin\nconfirma entrega

ENTREGADO --> [*]
CANCELADA --> [*]

note right of ENTREGADO
  Estado terminal.
  Habilita al cliente
  para reseñar productos.
end note

note left of CANCELADA
  No se puede cancelar
  desde ENVIADO ni ENTREGADO
  (el producto ya salió).
end note

note bottom
  <b>Casos que NO crean Order</b> (quedan en FailedCheckoutAttempt):
  • out_of_stock         · stock insuficiente al reservar
  • payment_rejected     · pasarela rechaza el pago
  • payment_unavailable  · Circuit Breaker abierto
  • inventory_unavailable· Inventory Service caído
end note

@enduml
"""


PLANTUML_COMPONENTES = r"""@startuml DiagramaComponentes
skinparam shadowing false
skinparam component {
  BackgroundColor #DCE6F1
  BorderColor #1F497D
  FontColor #172026
}
skinparam interface {
  BackgroundColor #FFFFFF
  BorderColor #1F497D
}
skinparam database {
  BackgroundColor #FEF3C7
  BorderColor #CA8A04
}
title <b>Diagrama de Componentes — Arquitectura de microservicios</b>

package "Cliente" {
  [Frontend SPA\nReact + Vite] as front
}

package "Capa de gateway" {
  [API Gateway\nNginx :80] as gw
}

package "Microservicios FastAPI" {
  [Auth Service\n:8001] as auth
  [Catalog Service\n:8002] as catalog
  [Inventory Service\n:8003] as inventory
  [Commerce Service\n:8004] as commerce
  [Payment Service\n:8005] as payment
}

package "Persistencia (MySQL 8.4 multi-schema)" {
  database "auth_db" as authdb
  database "catalog_db" as catdb
  database "inventory_db" as invdb
  database "commerce_db" as comdb
  database "payments_db" as paydb
}

package "Infraestructura compartida" {
  [Redis 7\nCache + Locks + CB] as redis
  [Mailhog\nSMTP local] as smtp
  [Payment Mock\nPasarela simulada] as mock
  [phpMyAdmin\nInspección DB] as pma
}

' ═══ Cliente → Gateway ═══
front --> gw : HTTPS\nJSON + JWT

' ═══ Gateway → Microservicios ═══
gw --> auth      : /api/auth, /api/users
gw --> catalog   : /api/products, /api/catalog
gw --> inventory : /api/inventory, /api/variants
gw --> commerce  : /api/cart, /api/checkout,\n/api/orders, /api/reviews
gw --> payment   : /api/payments

' ═══ Microservicios → su DB ═══
auth --> authdb
catalog --> catdb
inventory --> invdb
commerce --> comdb
payment --> paydb

' ═══ Microservicios → Redis ═══
catalog ..> redis : Cache-Aside
inventory ..> redis : Lock distribuido
payment ..> redis : Circuit Breaker state

' ═══ Microservicios → SMTP ═══
auth ..> smtp : welcome email
commerce ..> smtp : order emails

' ═══ Payment → Mock ═══
payment ..> mock : POST /charge

' ═══ Comunicación REST entre microservicios (SAGA) ═══
commerce ..> inventory : reserve / confirm / release\nvariants/by-ids
commerce ..> payment   : POST /payments (charge)
catalog ..> inventory  : stock-summary\nvariants by product
commerce ..> catalog   : PUT rating update

note bottom of commerce
  <b>Orquestador de la SAGA</b> orquestada
  síncrona REST. Garantiza consistencia
  eventual con compensaciones HTTP.
end note

@enduml
"""


PLANTUML_DESPLIEGUE = r"""@startuml DiagramaDespliegue
skinparam shadowing false
skinparam node {
  BackgroundColor #DCE6F1
  BorderColor #1F497D
  FontColor #172026
}
skinparam database {
  BackgroundColor #FEF3C7
  BorderColor #CA8A04
}
skinparam component {
  BackgroundColor #FFFFFF
  BorderColor #1F497D
}
title <b>Diagrama de Despliegue — Contenedores Docker</b>\n<i>12 contenedores en una red Docker bridge (tienda_net)</i>

node "<<device>>\nEquipo del estudiante\n(Windows + Docker Desktop)" as host {

  node "<<network>>\nDocker network: tienda_net" as net {

    node "tienda_gateway\n<<container>>\nimage: nginx:alpine\nport 80→80" as cgw {
      component "api-gateway\n(Nginx routing,\nCORS, rate-limit)" as gw
    }

    node "tienda_auth\n<<container>>\nbuild: auth-service\nport 8001→8001" as cauth {
      component "FastAPI\nAuth Service" as authsvc
    }

    node "tienda_catalog\n<<container>>\nbuild: catalog-service\nport 8002→8002" as ccat {
      component "FastAPI\nCatalog Service" as catsvc
    }

    node "tienda_inventory\n<<container>>\nbuild: inventory-service\nport 8003→8003" as cinv {
      component "FastAPI\nInventory Service" as invsvc
    }

    node "tienda_commerce\n<<container>>\nbuild: commerce-service\nport 8004→8004" as ccom {
      component "FastAPI\nCommerce Service" as comsvc
    }

    node "tienda_payment\n<<container>>\nbuild: payment-service\nport 8005→8005" as cpay {
      component "FastAPI\nPayment Service" as paysvc
    }

    node "tienda_payment_mock\n<<container>>\nport 9000→9000" as cmock {
      component "Pasarela simulada\n(APPROVED/REJECTED/\nPENDING/FAILED)" as mocksvc
    }

    node "tienda_digital_mysql\n<<container>>\nimage: mysql:8.4\nport 3306→3306" as cmysql {
      database "auth_db" as adb
      database "catalog_db" as cdb
      database "inventory_db" as idb
      database "commerce_db" as comdb
      database "payments_db" as pdb
    }

    node "tienda_redis\n<<container>>\nimage: redis:7\nport 6379→6379" as credis {
      component "Cache\nLocks\nCB state" as redissvc
    }

    node "tienda_mailhog\n<<container>>\nimage: mailhog\nport 1025 + 8025" as cmh {
      component "SMTP + UI web" as smtpsvc
    }

    node "tienda_phpmyadmin\n<<container>>\nport 8080→80" as cpma {
      component "phpMyAdmin\n(inspección DB)" as pmasvc
    }
  }
}

node "<<device>>\nCliente / Admin\n(Navegador en host)" as browser {
  component "http://localhost\n(puerto 80 del gateway)" as ui
}

node "<<device>>\nFrontend dev server\n(npm run dev)" as fronthost {
  component "React + Vite\npuerto 5173 (host)" as react
}

' Conexiones físicas/lógicas
browser --> gw : HTTPS / JSON
react --> gw   : HTTPS / JSON (dev)

gw --> authsvc
gw --> catsvc
gw --> invsvc
gw --> comsvc
gw --> paysvc

authsvc --> adb
catsvc --> cdb
invsvc --> idb
comsvc --> comdb
paysvc --> pdb

catsvc ..> redissvc
invsvc ..> redissvc
paysvc ..> redissvc

authsvc ..> smtpsvc
comsvc ..> smtpsvc

paysvc ..> mocksvc

note bottom of cmysql
  <b>Database per Service:</b> cada
  microservicio tiene su esquema
  aislado, con un usuario MySQL
  propio que solo accede a él.
  (Validado por Conformity Monkey.)
end note

@enduml
"""


# =========================================================================
# Construcción del documento
# =========================================================================

def construir_srs():
    """Construye y guarda el SRS consolidado con requisitos y diagramas."""
    doc = Document()
    set_doc_lang(doc); set_margins(doc)

    # =================================================================
    # PORTADA
    # =================================================================
    add_cover(doc)

    # =================================================================
    # RESUMEN EJECUTIVO
    # =================================================================
    add_h1(doc, "Resumen ejecutivo")
    add_para(doc,
        "Este documento define de forma clara, verificable y trazable los requisitos de una "
        "aplicación de tienda digital destinada a Distrito Urbano, una empresa de moda urbana "
        "con ventas en línea. La solución cubre dos frentes integrados: la experiencia de compra "
        "de los clientes y el panel administrativo del negocio. Desde el lado comercial, el sistema "
        "permite publicar productos con variantes de color y talla, vender en línea y procesar "
        "pagos a través de una pasarela externa. Desde el lado operativo, permite controlar "
        "inventario en tiempo real, pedidos con ciclo de vida formal, empleados, gastos, reseñas "
        "moderadas y reportes financieros con cálculo real de COGS y margen.")
    add_para(doc,
        "Esta revisión del SRS sustituye la línea base original (SRS-TD-EMP-001, abril 2026) y "
        "consolida los 9 requisitos funcionales (RF-01 a RF-09) y 3 requisitos no funcionales "
        "(RNF-01 a RNF-03) priorizados para el MVP. Cada requisito incluye su declaración, "
        "criterio de aceptación en formato de historia de usuario y tareas técnicas asociadas. "
        "El proyecto se implementa con arquitectura de microservicios FastAPI, base de datos "
        "MySQL multi-schema, frontend React y Docker Compose como plataforma local.")

    # Índice
    add_h2(doc, "Índice general")
    secs = [
        ("1.", "Introducción"),
        ("2.", "Descripción general"),
        ("3.", "Requisitos funcionales (RF-01 a RF-09)"),
        ("4.", "Requisitos no funcionales (RNF-01 a RNF-03)"),
        ("5.", "Restricciones del dominio"),
        ("6.", "Actores del sistema"),
        ("7.", "Diagrama de casos de uso"),
        ("8.", "Especificación de casos de uso (≥ 3 únicos del dominio)"),
        ("9.", "Diagramas de comportamiento UML (secuencia, actividad, estado)"),
        ("10.", "Diagramas estructurales UML (clases, componentes, despliegue)"),
        ("11.", "Reglas de negocio y consideraciones adicionales"),
        ("12.", "Índice de trazabilidad"),
    ]
    add_table(doc, ["No.", "Sección"], secs, widths_cm=[2, 14])

    doc.add_page_break()

    # =================================================================
    # 1. INTRODUCCIÓN
    # =================================================================
    add_h1(doc, "1. Introducción")

    add_h2(doc, "1.1 Propósito")
    add_para(doc,
        "El propósito de este documento es especificar los requisitos del software para una "
        "aplicación de tienda digital orientada a Distrito Urbano. El SRS actúa como contrato "
        "funcional y técnico entre los interesados del proyecto, proporciona una base para "
        "planificación, arquitectura, construcción, pruebas y aceptación, y reduce la ambigüedad "
        "al convertir necesidades del negocio en requisitos verificables.")
    add_para(doc,
        "Este documento no describe cómo se programará cada componente a nivel de diseño "
        "detallado; describe qué debe hacer el sistema, qué restricciones debe respetar y qué "
        "criterios permitirán validar su cumplimiento.")

    add_h2(doc, "1.2 Alcance del producto")
    add_para(doc,
        "La solución comprende una tienda digital de marca propia para Distrito Urbano y un panel "
        "administrativo asociado. El cliente final podrá navegar el catálogo, consultar productos "
        "con variantes (color y talla), agregar artículos al carrito, completar el checkout, pagar "
        "por una pasarela externa, consultar sus pedidos y dejar reseñas posteriores a la compra "
        "que pasan por moderación administrativa. El administrador podrá gestionar el catálogo, "
        "controlar el inventario por variante con alertas de stock mínimo, administrar pedidos a "
        "través de su ciclo logístico, registrar empleados, salarios, gastos y consultar reportes "
        "operativos y financieros con cálculo real de COGS y margen.")

    add_h2(doc, "1.3 Definiciones, siglas y acrónimos")
    add_table(doc, ["Término", "Definición"], [
        ("SRS",  "Software Requirements Specification; documento formal que define el comportamiento y las restricciones del sistema."),
        ("RF",   "Requisito funcional; capacidad observable que el sistema debe proporcionar."),
        ("RNF",  "Requisito no funcional; atributo de calidad o restricción del sistema."),
        ("UML",  "Unified Modeling Language; lenguaje estándar para modelar sistemas de software."),
        ("API",  "Interfaz de programación de aplicaciones para integración entre componentes."),
        ("SKU",  "Stock Keeping Unit; código único de inventario asociado a una variante."),
        ("Checkout", "Flujo de confirmación de compra, validación y pago del pedido."),
        ("COGS", "Cost of Goods Sold; costo de los productos efectivamente vendidos en un período."),
        ("SAGA", "Patrón de transacciones distribuidas con compensaciones. Se aplica orquestada y síncrona en el checkout."),
        ("JWT",  "JSON Web Token; mecanismo de autenticación basado en tokens firmados."),
        ("PlantUML", "Lenguaje textual para describir diagramas UML que pueden ser renderizados a imágenes."),
        ("MVP",  "Minimum Viable Product; versión mínima funcional del producto entregada al cierre del proyecto académico."),
    ], widths_cm=[3, 13])

    add_h2(doc, "1.4 Referencias")
    add_table(doc, ["Referencia", "Uso dentro del documento"], [
        ("ISO/IEC/IEEE 29148:2018", "Requisitos y documentación de ingeniería de requisitos."),
        ("ISO/IEC 25010:2023",      "Modelo de calidad de producto de software."),
        ("UML 2.5.1 / PlantUML",    "Notación de modelado utilizada en los diagramas embebidos en este SRS."),
        ("OWASP ASVS",              "Buenas prácticas de seguridad, autenticación y registro de eventos."),
        ("Guía Scrum 2020",         "Marco de trabajo utilizado para la gestión iterativa del proyecto."),
        ("Netflix Simian Army (2011)", "Inspiración para los experimentos de Chaos Engineering aplicados al MVP."),
    ], widths_cm=[6, 10])

    add_h2(doc, "1.5 Visión general del documento")
    add_para(doc,
        "Las secciones 2 a 4 describen el contexto del negocio, las funciones del producto y los "
        "requisitos funcionales y no funcionales. La sección 5 lista las restricciones del dominio. "
        "La sección 6 identifica los actores del sistema. La sección 7 presenta el diagrama de "
        "casos de uso (en código PlantUML embebido). La sección 8 contiene la especificación "
        "detallada de los casos de uso principales, todos ellos representativos de funciones "
        "únicas del dominio de Distrito Urbano. Las secciones 9 y 10 presentan los diagramas de "
        "comportamiento y estructurales en PlantUML. La sección 11 documenta reglas de negocio y "
        "consideraciones de arquitectura. La sección 12 cierra con la matriz de trazabilidad.")

    # =================================================================
    # 2. DESCRIPCIÓN GENERAL
    # =================================================================
    add_h1(doc, "2. Descripción general")

    add_h2(doc, "2.1 Contexto del negocio")
    add_para(doc,
        "Distrito Urbano requiere un canal digital propio que le permita vender sus productos en "
        "línea sin depender únicamente de canales manuales o redes sociales. El modelo operativo "
        "actual presenta riesgos comunes: poca visibilidad del inventario real, dificultad para "
        "controlar productos con variantes (color y talla), seguimiento manual de pedidos, baja "
        "trazabilidad de los pagos y ausencia de reportes financieros unificados. El sistema "
        "propuesto centraliza estas operaciones en una sola solución tecnológica con arquitectura "
        "moderna basada en microservicios.")

    add_h2(doc, "2.2 Objetivos de negocio")
    add_table(doc, ["ID", "Objetivo"], [
        ("OBJ-01", "Incrementar las ventas digitales de Distrito Urbano mediante un canal de compra en línea estable y usable."),
        ("OBJ-02", "Centralizar la gestión del catálogo, variantes (color + talla) y existencias en una sola plataforma."),
        ("OBJ-03", "Reducir errores operativos en pedidos, pagos y stock mediante la SAGA orquestada y la validación transaccional."),
        ("OBJ-04", "Proveer información financiera con COGS y márgenes reales para apoyar decisiones del negocio."),
        ("OBJ-05", "Mantener una experiencia de marca coherente entre la tienda digital pública y el panel administrativo."),
    ], widths_cm=[2.5, 13.5])

    add_h2(doc, "2.3 Stakeholders y usuarios objetivo")
    add_table(doc, ["Stakeholder / usuario", "Necesidad principal", "Prioridad"], [
        ("Dirección del negocio",
         "Patrocina la solución, define objetivos comerciales y lineamientos de marca.", "Alta"),
        ("Administrador del e-commerce",
         "Gestiona catálogo, inventario, pedidos, clientes, empleados, gastos y reportes.", "Alta"),
        ("Cliente final",
         "Navega el catálogo, compra productos, paga, consulta pedidos y deja reseñas.", "Alta"),
        ("Área contable/financiera",
         "Usa la información operativa para control interno y toma de decisiones.", "Media"),
        ("Pasarela de pago (sistema externo)",
         "Autoriza, rechaza o confirma transacciones monetarias mediante API externa.", "Alta"),
        ("Servicio SMTP (sistema externo)",
         "Entrega correos transaccionales: bienvenida, confirmación de pedido, cambios de estado.", "Media"),
        ("Equipo de tecnología (proyecto Scrum)",
         "Implementa, despliega, integra, monitorea y mantiene la solución.", "Alta"),
    ], widths_cm=[5, 9, 2])

    add_h2(doc, "2.4 Funciones principales del producto")
    funcs = [
        "Registro y autenticación de clientes con JWT.",
        "Visualización del catálogo con categorías, búsqueda, filtros y variantes (color + talla).",
        "Gestión de carrito y proceso de checkout con validación transaccional de inventario.",
        "Integración con pasarela de pago externa (simulada con 4 escenarios para el MVP).",
        "Creación y seguimiento de pedidos con timeline de estados visible al cliente.",
        "Reseñas y valoraciones con moderación administrativa antes de publicar.",
        "Panel administrativo para productos, variantes, stock, pedidos, empleados y reportes.",
        "Configuración visual de la tienda: logo, colores corporativos y datos informativos.",
        "Registro de bitácora con correlation-id para trazabilidad de operaciones críticas.",
        "Dashboard financiero con cálculo real de COGS, margen bruto, utilidad neta y exportación a PDF/CSV.",
    ]
    for f in funcs:
        add_bullet(doc, f)

    add_h2(doc, "2.5 Entorno operativo y tecnológico")
    add_table(doc, ["Elemento", "Descripción"], [
        ("Frontend cliente",  "React 18 + Vite, HTML5 y CSS3. Responsive para móvil y escritorio."),
        ("Panel administrativo", "Misma SPA React, con rutas protegidas por rol admin (JWT)."),
        ("Backend",            "Arquitectura de microservicios FastAPI (5 servicios: Auth, Catalog, Inventory, Commerce, Payment)."),
        ("API Gateway",        "Nginx en puerto 80. Routing por path, CORS, rate-limit y propagación de correlation-id."),
        ("Base de datos",      "MySQL 8.4 con Database per Service (5 esquemas aislados con GRANT por usuario)."),
        ("Cache y locks",      "Redis 7: Cache-Aside (catálogo), lock distribuido (inventario), estado del Circuit Breaker (pagos)."),
        ("Mensajería",         "Mailhog como servidor SMTP local para correos transaccionales (puerto 8025 UI)."),
        ("Pasarela de pago",   "Simulador local (payment-mock) con 4 escenarios: APPROVED, REJECTED, PENDING, FAILED."),
        ("Entorno de ejecución","Docker Compose con 12 contenedores en una red bridge tienda_net."),
    ], widths_cm=[4, 12])

    # =================================================================
    # 3. REQUISITOS FUNCIONALES
    # =================================================================
    add_h1(doc, "3. Requisitos funcionales")
    add_para(doc,
        "Todos los requisitos funcionales se redactan con identificador único y verbo obligatorio. "
        "Para esta línea base se priorizan 9 requisitos (RF-01 a RF-09) que cubren el alcance "
        "completo del MVP. Cada requisito incluye: declaración, actor principal, prioridad, "
        "criterio de aceptación en formato de historia de usuario y tareas (tasks) que materializan "
        "su implementación.")

    # Tabla resumen
    add_h2(doc, "3.1 Resumen de requisitos funcionales")
    add_table(doc,
        ["ID", "Título", "Actor", "Prioridad"],
        [(rf[0], rf[1], rf[3], rf[4]) for rf in RF_DATA],
        widths_cm=[1.5, 7, 4.5, 2])

    # Detalle de cada RF
    add_h2(doc, "3.2 Detalle de los requisitos funcionales")
    for rf_id, titulo, decl, actor, prio, ac, tasks in RF_DATA:
        add_h3(doc, f"{rf_id} — {titulo}")
        add_table(doc, ["Campo", "Detalle"], [
            ("ID",           rf_id),
            ("Título",       titulo),
            ("Declaración",  decl),
            ("Actor",        actor),
            ("Prioridad",    prio),
            ("Criterio de aceptación", ac),
        ], widths_cm=[3.5, 12.5])
        add_h4(doc, "Tareas asociadas")
        add_table(doc, ["#", "Descripción"],
                  [(str(i+1), t) for i, t in enumerate(tasks)],
                  widths_cm=[1, 15])

    # =================================================================
    # 4. REQUISITOS NO FUNCIONALES
    # =================================================================
    add_h1(doc, "4. Requisitos no funcionales")
    add_para(doc,
        "Los requisitos no funcionales describen las propiedades de calidad de la solución. Esta "
        "línea base prioriza 3 requisitos integrales (RNF-01 a RNF-03) que abarcan rendimiento, "
        "seguridad y aspectos transversales como usabilidad, respaldo y trazabilidad.")

    add_h2(doc, "4.1 Resumen de requisitos no funcionales")
    add_table(doc,
        ["ID", "Título", "Prioridad"],
        [(rnf[0], rnf[1], rnf[3]) for rnf in RNF_DATA],
        widths_cm=[2, 12, 2])

    add_h2(doc, "4.2 Detalle de los requisitos no funcionales")
    for rnf_id, titulo, decl, prio, ac, tasks in RNF_DATA:
        add_h3(doc, f"{rnf_id} — {titulo}")
        add_table(doc, ["Campo", "Detalle"], [
            ("ID",           rnf_id),
            ("Título",       titulo),
            ("Declaración",  decl),
            ("Prioridad",    prio),
            ("Criterio de aceptación", ac),
        ], widths_cm=[3.5, 12.5])
        add_h4(doc, "Tareas asociadas")
        add_table(doc, ["#", "Descripción"],
                  [(str(i+1), t) for i, t in enumerate(tasks)],
                  widths_cm=[1, 15])

    # =================================================================
    # 5. RESTRICCIONES DEL DOMINIO
    # =================================================================
    add_h1(doc, "5. Restricciones del dominio")
    add_para(doc,
        "Las restricciones que siguen son condiciones invariantes derivadas del dominio comercial "
        "y técnico de Distrito Urbano. Deben respetarse en todas las decisiones de diseño e "
        "implementación.")

    add_table(doc, ["ID", "Restricción"], [
        ("CT-01", "El sistema se diseña para una única empresa (Distrito Urbano); no se contempla multi-tenant ni marketplace abierto."),
        ("CT-02", "El backend debe implementarse en Python con FastAPI, separado en microservicios con responsabilidades bien definidas."),
        ("CT-03", "La base de datos debe ser MySQL 8.4 con Database per Service (cada microservicio tiene su esquema aislado y su usuario MySQL con GRANT exclusivo)."),
        ("CT-04", "El frontend debe construirse con React 18 + Vite, HTML5 y CSS3."),
        ("CT-05", "El procesamiento de pagos se delega a una pasarela externa simulada; el sistema NO almacena datos completos de tarjeta."),
        ("CT-06", "Solo se persiste un pedido (Order) cuando el checkout llega a PAID. Casos fallidos quedan en FailedCheckoutAttempt para auditoría."),
        ("CT-07", "Las reseñas solo se permiten sobre productos comprados y entregados. Toda reseña entra como pendiente y requiere aprobación admin antes de publicarse."),
        ("CT-08", "Cada variante de producto tiene SKU único global. La combinación (producto, color, talla) también es única."),
        ("CT-09", "Idioma de interfaz: español de Colombia (es-CO). Moneda única: peso colombiano (COP). Zona horaria: America/Bogota."),
        ("CT-10", "Únicamente existen dos roles: customer y admin. No hay roles intermedios."),
        ("CT-11", "Cada cliente autenticado posee un único carrito en estado abierto a la vez. No se admiten carritos anónimos."),
        ("CT-12", "La primera versión NO incluye ERP, contabilidad tributaria completa, facturación electrónica oficial ni logística avanzada."),
    ], widths_cm=[1.8, 14.2])

    # =================================================================
    # 6. ACTORES DEL SISTEMA
    # =================================================================
    add_h1(doc, "6. Actores del sistema")
    add_para(doc,
        "Se identifican los siguientes actores que interactúan con la plataforma. Los actores "
        "humanos representan personas que usan la interfaz; los actores externos representan "
        "sistemas con los que la plataforma se integra a través de APIs.")

    add_table(doc, ["Actor", "Tipo", "Responsabilidades principales"], [
        ("Cliente", "Humano (primario)",
         "Registrarse, iniciar sesión, explorar catálogo, buscar y filtrar productos, consultar "
         "detalle, gestionar carrito, realizar checkout, consultar pedidos, reseñar productos "
         "entregados y recibir notificaciones."),
        ("Administrador", "Humano (primario)",
         "Gestionar catálogo (categorías, productos, imágenes), administrar inventario "
         "(variantes color+talla, movimientos, alertas), operar pedidos (transicionar estados), "
         "aprobar reseñas, consultar finanzas y dashboards, configurar la tienda y consultar la "
         "bitácora de auditoría."),
        ("Pasarela de Pago", "Sistema externo",
         "Recibir solicitudes de cobro desde el Payment Service y devolver autorización, rechazo, "
         "pendiente o fallo de la transacción. Implementada como simulador (payment-mock) para el MVP."),
        ("Servicio SMTP", "Sistema externo",
         "Recibir y entregar correos transaccionales del sistema: bienvenida al registrarse, "
         "confirmación de pedido, cambios de estado y notificaciones administrativas. "
         "Implementado con Mailhog para el entorno local."),
        ("Sistema (procesos internos)", "Auto-actor",
         "Scheduler de Inventory que expira reservas vencidas; worker reconciler de Payment que "
         "reintenta pagos PENDING/FAILED; healthchecks que verifican dependencias."),
    ], widths_cm=[4, 3, 9])

    # =================================================================
    # 7. DIAGRAMA DE CASOS DE USO
    # =================================================================
    add_h1(doc, "7. Diagrama de casos de uso")
    add_para(doc,
        "El diagrama UML de casos de uso muestra los actores que interactúan con el sistema y las "
        "principales capacidades en infinitivo. Se modelan dos tipos de relaciones: "
        "<<include>> indica que un caso de uso destino se ejecuta obligatoriamente como parte del "
        "caso origen (por ejemplo, Realizar checkout INCLUYE Procesar pago); "
        "<<extend>> indica que un caso destino extiende opcionalmente el comportamiento del caso "
        "origen bajo ciertas condiciones (por ejemplo, Buscar producto EXTIENDE Explorar catálogo).")
    add_para(doc,
        "El siguiente bloque contiene el código PlantUML del diagrama. Para renderizarlo, copiar el "
        "código y pegarlo en https://www.plantuml.com/plantuml/uml/ o utilizar la extensión "
        "PlantUML disponible en VSCode, IntelliJ y la mayoría de IDEs.",
        italic=True)

    add_code_block(doc, PLANTUML_CASOS_USO,
                   title="Diagrama de Casos de Uso", lang="plantuml")

    add_h2(doc, "7.1 Lectura del diagrama")
    add_para(doc,
        "El cliente concentra los casos de uso del flujo comercial: navegación del catálogo, "
        "construcción del carrito, ejecución del checkout, seguimiento de pedidos y publicación "
        "de reseñas. El administrador concentra los casos de uso operativos: gestión de catálogo, "
        "configuración de variantes Nike-style, operación logística de pedidos, moderación de "
        "reseñas, consulta financiera y configuración de la tienda. La pasarela de pago aparece "
        "como sistema externo vinculado al caso de uso Procesar pago. El servicio SMTP aparece "
        "como sistema externo vinculado a Recibir notificación y a Operar pedidos (porque el "
        "cambio de estado dispara un correo al cliente).")

    # =================================================================
    # 8. ESPECIFICACIÓN DE CASOS DE USO
    # =================================================================
    add_h1(doc, "8. Especificación de casos de uso")
    add_para(doc,
        "Se detallan a continuación los casos de uso principales del sistema. Cada uno incluye: "
        "actor principal, requisitos funcionales asociados, objetivo, precondición, postcondición "
        "(éxito y fallo), flujo principal paso a paso y flujos alternativos (caminos no-felices). "
        "Los casos de uso seleccionados representan funciones únicas del dominio de Distrito Urbano: "
        "la SAGA orquestada del checkout, la moderación de reseñas validadas por compra, la "
        "configuración de variantes color+talla con stock por combinación, la transición logística "
        "del pedido con notificación, y la consulta financiera con COGS real. NO se incluyen casos "
        "genéricos (como 'iniciar sesión') porque su comportamiento no difiere del estándar de "
        "cualquier aplicación web autenticada.")

    add_h2(doc, "8.1 Catálogo de casos de uso especificados")
    add_table(doc, ["ID", "Caso de uso", "Actor principal", "RF asociado"],
              [(cu["id"], cu["nombre"], cu["actor"], cu["rf"]) for cu in CASOS_DE_USO],
              widths_cm=[1.5, 8, 4, 2.5])

    # Detalle de cada caso
    for i, cu in enumerate(CASOS_DE_USO, 1):
        add_h2(doc, f"8.{i+1} {cu['id']} — {cu['nombre']}")
        add_table(doc, ["Campo", "Detalle"], [
            ("Identificador",           cu["id"]),
            ("Nombre",                  cu["nombre"]),
            ("Actor principal",         cu["actor"]),
            ("RF asociado",             cu["rf"]),
            ("Objetivo",                cu["objetivo"]),
            ("Precondición",            cu["precondicion"]),
            ("Postcondición (éxito)",   cu["postcondicion_exito"]),
            ("Postcondición (fallo)",   cu["postcondicion_fallo"]),
        ], widths_cm=[3.5, 12.5])

        add_h4(doc, "Flujo principal (camino feliz)")
        for j, paso in enumerate(cu["flujo"], 1):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.6)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(f"{j}. "); set_run(run, bold=True, size=10.5, color=BLUE_HEAD)
            run = p.add_run(paso); set_run(run, size=10.5)

        if cu["alternativos"]:
            add_h4(doc, "Flujos alternativos (caminos no-felices)")
            for fa_id, fa_titulo, fa_pasos in cu["alternativos"]:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(2)
                run = p.add_run(f"{fa_id} — {fa_titulo}")
                set_run(run, bold=True, size=11, color=ACCENT)

                for j, paso in enumerate(fa_pasos, 1):
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Cm(1.2)
                    p.paragraph_format.space_after = Pt(2)
                    run = p.add_run(f"{j}. "); set_run(run, bold=True, size=10.5, color=ACCENT)
                    run = p.add_run(paso); set_run(run, size=10.5)

    # =================================================================
    # 9. DIAGRAMAS DE COMPORTAMIENTO
    # =================================================================
    add_h1(doc, "9. Diagramas de comportamiento UML")
    add_para(doc,
        "Los diagramas de comportamiento describen cómo se comportan los objetos del sistema a lo "
        "largo del tiempo. Se incluyen tres diagramas representativos: secuencia del checkout "
        "(coordinación SAGA), actividad del proceso de compra del cliente (flujo end-to-end) y "
        "estado del ciclo de vida del pedido. Todos los códigos PlantUML siguientes son completos "
        "y renderizables sin modificaciones.")

    add_h2(doc, "9.1 Diagrama de secuencia — Realizar checkout (SAGA)")
    add_para(doc,
        "Este diagrama detalla la interacción temporal entre el cliente, el frontend, el gateway, "
        "los microservicios Commerce, Inventory y Payment, y la pasarela mock durante el flujo "
        "más crítico del sistema. Muestra cómo Commerce orquesta la SAGA: reserve → charge → "
        "confirm, con compensaciones HTTP si cualquier paso falla.")

    add_code_block(doc, PLANTUML_SECUENCIA,
                   title="Diagrama de Secuencia — Checkout SAGA", lang="plantuml")

    add_h2(doc, "9.2 Diagrama de actividad — Proceso de compra del cliente")
    add_para(doc,
        "Este diagrama representa el flujo de trabajo del cliente desde que inicia sesión hasta "
        "que recibe la confirmación de la compra, incluyendo los puntos de decisión: ¿hay stock? "
        "y ¿SAGA OK? Cada rama no-feliz tiene su tratamiento explícito.")

    add_code_block(doc, PLANTUML_ACTIVIDAD,
                   title="Diagrama de Actividad — Proceso de compra", lang="plantuml")

    add_h2(doc, "9.3 Diagrama de estado — Ciclo de vida del Pedido")
    add_para(doc,
        "Este diagrama representa el ciclo de vida de la entidad Order desde su creación hasta su "
        "estado terminal. Bajo la política MVP, un pedido solo existe en la base si el checkout "
        "termina exitosamente en estado PAID. A partir de ahí, el administrador puede transicionarlo "
        "a EN_PREPARACION, ENVIADO y finalmente ENTREGADO. Desde PAID o EN_PREPARACION es posible "
        "cancelar. Los estados ENVIADO y ENTREGADO no admiten cancelación.")

    add_code_block(doc, PLANTUML_ESTADO,
                   title="Diagrama de Estado — Ciclo de vida del Pedido", lang="plantuml")

    # =================================================================
    # 10. DIAGRAMAS ESTRUCTURALES
    # =================================================================
    add_h1(doc, "10. Diagramas estructurales UML")
    add_para(doc,
        "Los diagramas estructurales describen las partes que componen el sistema y cómo se "
        "relacionan, independientemente del tiempo. Se incluyen el diagrama de clases del modelo "
        "de dominio, el diagrama de componentes (arquitectura de microservicios) y el diagrama "
        "de despliegue (contenedores Docker).")

    add_h2(doc, "10.1 Diagrama de clases — Modelo de dominio")
    add_para(doc,
        "El diagrama de clases muestra las entidades persistentes principales agrupadas por "
        "microservicio. Cada microservicio posee su propio esquema MySQL aislado, por lo que las "
        "asociaciones entre entidades de distintos microservicios son referencias lógicas "
        "(estereotipo <<ref-logica>>) y NO claves foráneas físicas. Esto permite escalar y "
        "desplegar cada microservicio de forma independiente.")

    add_code_block(doc, PLANTUML_CLASES,
                   title="Diagrama de Clases — Modelo de dominio", lang="plantuml")

    add_h2(doc, "10.2 Diagrama de componentes — Arquitectura de microservicios")
    add_para(doc,
        "Este diagrama presenta los componentes lógicos del sistema y sus interfaces. La aplicación "
        "se compone del frontend SPA, el API Gateway en Nginx como punto único de entrada, los "
        "cinco microservicios FastAPI con responsabilidades bien definidas, la infraestructura "
        "compartida (MySQL multi-schema, Redis, Mailhog, payment-mock) y la herramienta auxiliar "
        "phpMyAdmin.")

    add_code_block(doc, PLANTUML_COMPONENTES,
                   title="Diagrama de Componentes", lang="plantuml")

    add_h2(doc, "10.3 Diagrama de despliegue — Contenedores Docker")
    add_para(doc,
        "Este diagrama muestra cómo se distribuyen los componentes en el entorno local. Toda la "
        "aplicación corre sobre Docker Compose, organizada en 12 contenedores conectados a una "
        "red bridge llamada tienda_net. Cada contenedor expone su puerto al host para facilitar el "
        "debug, pero las comunicaciones internas usan los nombres de contenedor en la red Docker.")

    add_code_block(doc, PLANTUML_DESPLIEGUE,
                   title="Diagrama de Despliegue", lang="plantuml")

    # =================================================================
    # 11. REGLAS DE NEGOCIO Y CONSIDERACIONES
    # =================================================================
    add_h1(doc, "11. Reglas de negocio y consideraciones adicionales")

    add_h2(doc, "11.1 Reglas de negocio")
    add_table(doc, ["ID", "Regla de negocio"], [
        ("BR-01", "La compra se realiza siempre contra el catálogo de Distrito Urbano (una sola empresa)."),
        ("BR-02", "El descuento de stock ocurre a nivel de variante (color + talla), no a nivel de producto general."),
        ("BR-03", "Un pedido solo se considera pagado cuando la pasarela confirma la transacción con respuesta válida."),
        ("BR-04", "Las reseñas se habilitan únicamente para pedidos entregados y deben ser aprobadas por el administrador antes de publicarse."),
        ("BR-05", "Los gastos y salarios afectan los indicadores financieros del período seleccionado, no del período en que ocurrió la transacción."),
        ("BR-06", "Los cambios manuales de inventario exigen motivo y usuario responsable; quedan en StockMovement con correlation-id."),
        ("BR-07", "La SAGA del checkout es atómica desde la perspectiva del cliente: o se completa todo (PAID + stock descontado) o nada (carrito intacto)."),
        ("BR-08", "Los pedidos en estado ENVIADO o ENTREGADO no pueden ser cancelados: el producto ya salió del control de la tienda."),
        ("BR-09", "El COGS se calcula usando el unit_cost snapshot al momento del checkout, no el costo actual de la variante."),
        ("BR-10", "El Circuit Breaker de Payment se abre tras 5 fallos consecutivos en 60 segundos y se recupera tras 60 segundos en estado HALF_OPEN."),
    ], widths_cm=[1.8, 14.2])

    add_h2(doc, "11.2 Reglas de cálculo financiero")
    add_para(doc,
        "La fórmula de utilidad neta utilizada en el dashboard financiero es:")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("utilidad_neta = ventas_brutas - COGS - gastos_operativos - nómina")
    set_run(run, bold=True, italic=True, size=11, color=BLUE_HEAD)
    add_para(doc,
        "donde:",
        before=2)
    add_bullet(doc, "ventas_brutas = SUM(Order.total) en estados PAID, EN_PREPARACION, ENVIADO, ENTREGADO.")
    add_bullet(doc, "COGS = SUM(OrderItem.unit_cost × quantity) — usa snapshot del costo capturado en el checkout.")
    add_bullet(doc, "gastos_operativos = SUM(Expense.amount) en el período seleccionado.")
    add_bullet(doc, "nómina = SUM(Employee.salary) donde employment_status = 'active'.")
    add_bullet(doc, "margen_bruto = ventas_brutas - COGS")
    add_bullet(doc, "margen_bruto_% = margen_bruto / ventas_brutas × 100")
    add_bullet(doc, "margen_neto_% = utilidad_neta / ventas_brutas × 100")
    add_bullet(doc, "ticket_promedio = ventas_brutas / orders_count")

    add_h2(doc, "11.3 Patrones arquitecturales aplicados")
    add_table(doc, ["Patrón", "Aplicación"], [
        ("API Gateway",
         "Nginx en puerto 80. Routing por path, CORS, headers de seguridad, rate-limit en /auth/login, propagación de correlation-id."),
        ("SSO con JWT compartido",
         "El secreto HS256 se comparte entre los servicios; ningún servicio llama de vuelta a Auth para validar tokens."),
        ("Database per Service",
         "5 esquemas MySQL aislados + un usuario MySQL por servicio con GRANT exclusivo."),
        ("Cache-Aside",
         "Catalog cachea categorías, productos y settings en Redis con TTL 60–300 s. Si Redis cae, se sirve desde MySQL."),
        ("Lock distribuido",
         "Inventory adquiere SET NX EX en Redis por variante + SELECT FOR UPDATE en MySQL antes de reservar."),
        ("SAGA orquestada síncrona",
         "Commerce coordina reserve → charge → confirm/release como transacción distribuida sin commit atómico."),
        ("Circuit Breaker",
         "Payment rechaza requests cuando hay 5 fallos en 60 s; estado CLOSED/OPEN/HALF_OPEN persistido en Redis."),
        ("Retry con backoff exponencial",
         "Payment reintenta errores transitorios con 250 ms / 500 ms / 1 s. Distingue REJECTED (negocio) de errores de infraestructura."),
        ("Worker reconciler",
         "Payment ejecuta una tarea asíncrona cada 5 min que reintenta los pagos en estado PENDING o FAILED."),
        ("Idempotencia",
         "Commerce acepta Idempotency-Key en /checkout y evita doble cobro si el cliente reintenta con la misma clave."),
        ("Bitácora con correlation-id",
         "El gateway inyecta X-Correlation-Id en cada request y los servicios lo propagan a logs y a tablas de auditoría."),
    ], widths_cm=[5, 11])

    add_h2(doc, "11.4 Alcance fuera de esta línea base")
    add_para(doc,
        "Quedan explícitamente fuera del alcance: marketplace multiempresa, administración de varias "
        "tiendas, integración con facturación electrónica oficial, logística avanzada de última "
        "milla, CRM avanzado, recomendaciones basadas en inteligencia artificial y contabilidad "
        "legal completa.")

    add_h2(doc, "11.5 Riesgos relevantes del proyecto")
    add_table(doc, ["ID", "Riesgo", "Impacto / control"], [
        ("R-01", "Dependencia de la pasarela de pago",
         "Un fallo externo puede impedir confirmaciones de compra; se mitiga con el Circuit Breaker, "
         "reintentos exponenciales y el worker reconciler. El cliente nunca se queda con doble cobro."),
        ("R-02", "Concurrencia sobre inventario",
         "Compras simultáneas pueden generar sobreventa; se mitiga con lock distribuido Redis "
         "por variante + SELECT FOR UPDATE en MySQL al reservar."),
        ("R-03", "Sobrecarga funcional del MVP",
         "Agregar demasiados módulos puede comprometer calidad; se mitiga con priorización Scrum "
         "y entregas incrementales por sprint."),
        ("R-04", "Definición financiera insuficiente",
         "La utilidad neta puede variar si no se cierran criterios de costos y gastos; se mitiga "
         "con el snapshot de unit_cost en OrderItem y la validación con el área contable."),
        ("R-05", "Caída del servicio SMTP",
         "Si Mailhog cae, los correos transaccionales no se envían pero las acciones críticas "
         "(pedidos, cambios de estado) SÍ se persisten y notifican in-app."),
    ], widths_cm=[1.5, 5, 9.5])

    # =================================================================
    # 12. ÍNDICE DE TRAZABILIDAD
    # =================================================================
    add_h1(doc, "12. Índice de trazabilidad")
    add_para(doc,
        "La siguiente matriz relaciona los objetivos del negocio con los requisitos, casos de uso "
        "y artefactos UML asociados. Su finalidad es apoyar la validación, la priorización del "
        "backlog, la planificación de pruebas y el control de cambios del producto.")

    add_table(doc, ["Objetivo", "Descripción", "RF", "RNF", "Casos de uso", "Diagramas UML"], [
        ("OBJ-01", "Canal de venta digital",
         "RF-01 a RF-06", "RNF-01, RNF-02, RNF-03",
         "CU-01", "Casos uso, Secuencia, Actividad, Estado"),
        ("OBJ-02", "Control de catálogo e inventario",
         "RF-02, RF-03, RF-07", "RNF-02, RNF-03",
         "CU-03", "Casos uso, Clases, Componentes"),
        ("OBJ-03", "Operación de pedidos",
         "RF-06, RF-07", "RNF-02, RNF-03",
         "CU-04", "Casos uso, Estado, Secuencia"),
        ("OBJ-04", "Gestión financiera",
         "RF-08", "RNF-03",
         "CU-05", "Clases, Componentes"),
        ("OBJ-05", "Identidad y configuración",
         "RF-02, RF-09", "RNF-03",
         "CU-02", "Casos uso, Clases"),
    ], widths_cm=[1.6, 3.8, 2.8, 2.8, 2, 4])

    add_h2(doc, "Cierre del documento")
    add_para(doc,
        "Este SRS constituye la línea base actualizada de requisitos para Distrito Urbano. "
        "Cualquier cambio futuro deberá evaluarse respecto de impacto funcional, técnico, "
        "operativo y de negocio, y deberá reflejarse en versiones controladas del documento. La "
        "versión actual (SRS-TD-EMP-002) reemplaza a la línea base original de abril de 2026 "
        "consolidando 9 requisitos funcionales y 3 requisitos no funcionales priorizados para la "
        "entrega del MVP académico bajo metodología Scrum.")

    # Guardar
    out_path = BASE / "SRS_Tienda_Digital.docx"
    doc.save(out_path)
    print(f"OK  {out_path.name} generado ({out_path.stat().st_size // 1024} KB)")


if __name__ == "__main__":
    print("Generando SRS único...")
    construir_srs()
    print("Listo. Documento en docs/entregables/SRS_Tienda_Digital.docx")
